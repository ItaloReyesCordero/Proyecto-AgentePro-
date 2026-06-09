# -*- coding: utf-8 -*-
"""Genera los documentos formales (.docx) para el registro de obra ante INDECOPI:
   - Memoria Descriptiva
   - Manual de Usuario
Ejecutar con: python _gen_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

AZUL = RGBColor(0x1D, 0x4E, 0xD8)
GRIS = RGBColor(0x55, 0x55, 0x55)
NEGRO = RGBColor(0x11, 0x11, 0x11)

AUTOR = "Italo Eduardo Reyes Cordero"
CORREO = "italoreyescordero1@gmail.com"
DNI = "75220834"
DOMICILIO = "Jr. Grau N.º 419, distrito de Jauja, provincia de Jauja, departamento de Junín"
TELEFONO = "916085873"

# Coautores (estudiantes de la Universidad Continental) y docente asesor
COAUTOR2 = "Jack Joshua Bendezu Lagos"
DNI2 = "73940475"
COAUTOR3 = "Dickmar Wilber Julca Laureano"
DNI3 = "73086197"
ASESOR = "Maglioni Arana Caparachin"
DNI_ASESOR = "20038141"
UNIVERSIDAD = "Universidad Continental"

CIUDAD = "Huancayo, Perú"
FECHA_FIRMA = "Huancayo, 5 de junio de 2026"
ANIO = "2026"


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.font.color.rgb = NEGRO
    # Fuente para caracteres asiáticos/otros, evita que Word cambie acentos
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_title_page(doc, titulo, subtitulo):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AGENTEPRO")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Plataforma SaaS de automatización de atención al cliente con inteligencia artificial")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GRIS

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitulo)
    r.font.size = Pt(12)
    r.font.color.rgb = GRIS

    for _ in range(6):
        doc.add_paragraph()

    datos = [
        ("Autor principal y titular", AUTOR),
        ("Coautores", f"{COAUTOR2}; {COAUTOR3}"),
        ("Docente asesor", ASESOR),
        ("Institución", UNIVERSIDAD),
        ("Tipo de obra", "Programa de ordenador (software)"),
        ("Año de creación", ANIO),
        ("Lugar", CIUDAD),
        ("Correo de contacto", CORREO),
    ]
    for k, v in datos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{k}: ")
        r.bold = True
        p.add_run(v)

    doc.add_page_break()


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(15)
    r.bold = True
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = NEGRO
    r.font.size = Pt(12.5)
    r.bold = True
    return p


def para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullets(doc, items):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(it[0] + ": ")
            r.bold = True
            p.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def kv_table(doc, rows, headers=("Campo", "Detalle")):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    hdr[0].paragraphs[0].add_run(headers[0]).bold = True
    hdr[1].paragraphs[0].add_run(headers[1]).bold = True
    for a, b in rows:
        c = t.add_row().cells
        c[0].paragraphs[0].add_run(str(a)).bold = True
        c[1].paragraphs[0].add_run(str(b))
    doc.add_paragraph()
    return t


# ════════════════════════════════════════════════════════════════════════════
# 1) MEMORIA DESCRIPTIVA
# ════════════════════════════════════════════════════════════════════════════
def memoria_descriptiva(path):
    doc = Document()
    base_style(doc)
    add_title_page(doc, "Memoria Descriptiva",
                   "Descripción técnica de la obra · Programa de ordenador")

    h1(doc, "1. Presentación")
    para(doc,
         "El presente documento constituye la memoria descriptiva de la obra de software "
         "denominada «AgentePro», elaborada para acompañar su registro como programa de "
         "ordenador ante la Dirección de Derecho de Autor (DDA) del Instituto Nacional de Defensa "
         "de la Competencia y de la Protección de la Propiedad Intelectual (INDECOPI), al amparo "
         "del Decreto Legislativo N.º 822, Ley sobre el Derecho de Autor, y la Decisión Andina 351.")
    para(doc,
         "Su finalidad es explicar, en lenguaje claro y preciso, qué hace el programa, cuáles son "
         "sus funciones, su arquitectura y los lenguajes de programación empleados en su desarrollo.")

    h1(doc, "2. Identificación de la obra")
    kv_table(doc, [
        ("Título de la obra", "AgentePro"),
        ("Tipo de obra", "Programa de ordenador (software)"),
        ("Clase", "Obra literaria (software), según D. Leg. N.º 822"),
        ("Naturaleza", "Obra originaria e individual"),
        ("Año de creación", ANIO),
        ("País de origen", "Perú"),
        ("Lenguajes de programación", "Python, TypeScript, JavaScript y SQL"),
    ])

    h1(doc, "3. Autoría, datos del autor y titular")
    para(doc,
         f"AgentePro es una obra de software desarrollada en colaboración por estudiantes de la "
         f"{UNIVERSIDAD}, bajo la asesoría del docente {ASESOR}. La autoría principal y la titularidad "
         f"de los derechos patrimoniales corresponden a {AUTOR}, quien concibió, dirigió y programó la "
         f"mayor parte de la obra. Participaron como coautores, con aportes complementarios, los "
         f"estudiantes {COAUTOR2} y {COAUTOR3}.")
    kv_table(doc, [
        ("Autor principal y titular", f"{AUTOR} — DNI N.º {DNI}"),
        ("Coautor", f"{COAUTOR2} — DNI N.º {DNI2} (estudiante de la {UNIVERSIDAD})"),
        ("Coautor", f"{COAUTOR3} — DNI N.º {DNI3} (estudiante de la {UNIVERSIDAD})"),
        ("Docente asesor", f"{ASESOR} — DNI N.º {DNI_ASESOR} (no es autor de la obra)"),
    ], headers=("Rol", "Identificación"))
    para(doc,
         "Los coautores han cedido sus derechos patrimoniales a favor del autor principal, de modo que "
         "la titularidad de dichos derechos recae de forma única y exclusiva en este último, "
         "conservando cada coautor sus derechos morales.")
    kv_table(doc, [
        ("Titular de los derechos patrimoniales", AUTOR),
        ("Nacionalidad", "Peruana"),
        ("Documento de identidad", f"DNI N.º {DNI}"),
        ("Domicilio", DOMICILIO),
        ("Correo electrónico", CORREO),
        ("Teléfono", TELEFONO),
    ])

    h1(doc, "4. Descripción general de la obra")
    para(doc,
         "AgentePro es una plataforma de software como servicio (SaaS) multiempresa, diseñada "
         "para que negocios del Perú automaticen la atención a sus clientes y la gestión comercial "
         "mediante inteligencia artificial. Cada negocio cliente (denominado «tenant» o inquilino) "
         "obtiene, de forma automática, un conjunto de herramientas que operan de manera continua, "
         "las veinticuatro horas del día, sin intervención manual.")
    para(doc,
         "El sistema centraliza en un único panel de control la comunicación por mensajería "
         "(WhatsApp e Instagram), la atención telefónica por voz, la administración de contactos y "
         "conversaciones, la generación de contenido para redes sociales y el seguimiento de "
         "clientes potenciales, todo ello sincronizado y visible en tiempo real.")

    h1(doc, "5. Problema que resuelve y objetivo")
    para(doc,
         "Muchos negocios pequeños y medianos pierden ventas porque no logran responder a tiempo a "
         "sus clientes, sobre todo fuera del horario laboral. Contratar personal de atención "
         "permanente resulta costoso. AgentePro resuelve este problema dotando a cada negocio "
         "de un asistente conversacional con inteligencia artificial que responde de inmediato, "
         "califica el interés del cliente, agenda citas y realiza seguimientos automáticos.")
    para(doc,
         "El objetivo de la obra es ofrecer una solución asequible, fácil de configurar y operable "
         "por personas sin conocimientos técnicos, que permita a cualquier negocio atender a sus "
         "clientes de manera profesional y permanente.")

    h1(doc, "6. Funcionalidades principales")
    bullets(doc, [
        ("Agente de WhatsApp con IA",
         "responde mensajes de los clientes de forma automática y natural, las veinticuatro horas, "
         "califica al cliente potencial y agenda citas."),
        ("Agente de voz con IA",
         "contesta y realiza llamadas telefónicas en español, conversando de manera natural con el "
         "cliente."),
        ("Gestión de contactos y conversaciones (CRM)",
         "registra automáticamente cada contacto y cada conversación, con sincronización opcional "
         "hacia plataformas de CRM externas."),
        ("Generación de contenido para Instagram",
         "crea publicaciones con texto e imagen mediante inteligencia artificial, listas para que "
         "el negocio las apruebe y publique."),
        ("Automatizaciones",
         "ejecuta seguimientos, recordatorios y reportes de manera programada y sin intervención "
         "humana."),
        ("Panel de control (dashboard)",
         "muestra en tiempo real las métricas del negocio: mensajes, llamadas, contactos y "
         "embudo de clientes potenciales."),
        ("Panel de superadministración",
         "permite al titular de la plataforma crear y administrar negocios, planes, cobros y "
         "soporte."),
        ("Administración de cobros",
         "gestiona el periodo de prueba gratuito y el cobro mensual por adelantado, con suspensión "
         "automática del servicio en caso de falta de pago."),
        ("Aislamiento de datos entre negocios",
         "garantiza que la información de cada negocio sea privada y no accesible por otros, "
         "reforzado a nivel de la capa de datos."),
    ])

    h1(doc, "7. Arquitectura y componentes")
    h2(doc, "7.1. Componente de servidor (backend)")
    para(doc,
         "Implementa la lógica de negocio, la seguridad, el aislamiento de datos entre empresas, la "
         "comunicación en tiempo real y la integración con servicios externos (modelos de "
         "inteligencia artificial, mensajería, telefonía, almacenamiento y CRM). Expone una interfaz "
         "de programación de aplicaciones (API) y procesa los eventos entrantes mediante webhooks.")
    h2(doc, "7.2. Componente de interfaz (frontend)")
    para(doc,
         "Es la aplicación web con la que interactúan los usuarios. Comprende la página pública de "
         "presentación, el panel del negocio y el panel de superadministración. Ofrece modo claro y "
         "oscuro y permite a cada negocio personalizar el color de su panel.")
    h2(doc, "7.3. Base de datos")
    para(doc,
         "Almacena de forma estructurada toda la información del sistema (negocios, usuarios, "
         "suscripciones, contactos, conversaciones, mensajes, llamadas, configuraciones y "
         "registros), con un esquema versionado mediante migraciones.")
    h2(doc, "7.4. Infraestructura de ejecución")
    para(doc,
         "La obra se distribuye en contenedores que incluyen el servidor de aplicaciones, los "
         "procesos de tareas en segundo plano, el servidor web con cifrado HTTPS, la base de datos, "
         "la memoria caché y un servicio de copias de seguridad automáticas.")

    h1(doc, "8. Lenguajes de programación y tecnologías")
    kv_table(doc, [
        ("Backend", "Python con el framework FastAPI; comunicación en tiempo real con Socket.IO; "
                    "acceso a datos con SQLAlchemy; tareas en segundo plano con Celery."),
        ("Frontend", "TypeScript y JavaScript con la biblioteca React y el empaquetador Vite; "
                     "estilos con Tailwind CSS."),
        ("Base de datos", "PostgreSQL; lenguaje SQL; migraciones con Alembic."),
        ("Infraestructura", "Docker y Docker Compose; servidor web Caddy con HTTPS automático; "
                            "memoria caché con Redis."),
        ("Inteligencia artificial", "Modelos de lenguaje de gran tamaño de Anthropic, "
                                    "orquestados por el sistema."),
    ], headers=("Capa", "Lenguajes y tecnologías"))

    h1(doc, "9. Módulos y modelos de datos principales")
    para(doc, "La obra organiza la información en los siguientes modelos de datos principales:")
    bullets(doc, [
        "Negocio (tenant) y usuario.",
        "Suscripción y configuración de cobros.",
        "Contacto, conversación y mensaje.",
        "Llamada y resumen de llamada.",
        "Configuración del agente de texto y del agente de voz.",
        "Publicación de Instagram.",
        "Automatización y ejecución de automatización.",
        "Registro de sincronización con CRM y registro de webhooks.",
        "Solicitud de recuperación de contraseña.",
    ])

    h1(doc, "10. Originalidad de la obra")
    para(doc,
         "El elemento original y distintivo de la obra reside en su arquitectura multiempresa con "
         "aislamiento de datos forzado a nivel de aplicación, en el aprovisionamiento automático de "
         "los canales de comunicación de cada negocio (mensajería, voz y redes sociales) y en la "
         "orquestación de un motor de inteligencia artificial para responder, calificar y dar "
         "seguimiento a los clientes de cada negocio de forma autónoma. La selección, disposición y "
         "organización de estos componentes, así como el código fuente que los implementa, "
         "constituyen una creación intelectual propia de los autores.")

    h1(doc, "11. Magnitud de la obra")
    kv_table(doc, [
        ("Backend", "Python / FastAPI / SQLAlchemy / Socket.IO — aproximadamente 130 archivos."),
        ("Frontend", "React / TypeScript / Vite — aproximadamente 60 archivos."),
        ("Base de datos", "PostgreSQL, con alrededor de diecisiete modelos de datos y sus migraciones."),
        ("Volumen total estimado", "Aproximadamente 15 000 líneas de código fuente."),
    ], headers=("Componente", "Detalle"))

    h1(doc, "12. Declaración")
    para(doc,
         f"Los autores declaran que la obra descrita en el presente documento es de su autoría, "
         f"original, y que no infringe derechos de terceros. El autor principal, {AUTOR}, es el titular "
         f"de los derechos patrimoniales de la obra. El derecho de autor sobre la obra nace con su "
         f"creación; el registro ante INDECOPI tiene por objeto constituir un medio de prueba de la "
         f"titularidad y de la fecha cierta de la creación.")
    doc.add_paragraph()
    doc.add_paragraph()
    firmantes = [
        (AUTOR, "Autor principal y titular de los derechos", DNI),
        (COAUTOR2, "Coautor", DNI2),
        (COAUTOR3, "Coautor", DNI3),
    ]
    for nombre, cargo, dni in firmantes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("_______________________________\n").bold = True
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(nombre + "\n").bold = True
        p2.add_run(cargo + "\n")
        p2.add_run(f"DNI N.º {dni}")
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(FECHA_FIRMA)
    r.italic = True
    r.font.color.rgb = GRIS

    doc.save(path)
    print("OK:", path)


# ════════════════════════════════════════════════════════════════════════════
# 2) MANUAL DE USUARIO
# ════════════════════════════════════════════════════════════════════════════
def manual_usuario(path):
    doc = Document()
    base_style(doc)
    add_title_page(doc, "Manual de Usuario",
                   "Guía de instalación, configuración y operación")

    h1(doc, "1. Introducción")
    para(doc,
         "AgentePro es una plataforma web que permite a un negocio automatizar la atención a "
         "sus clientes mediante inteligencia artificial: responde mensajes de WhatsApp e Instagram, "
         "atiende llamadas de voz, registra contactos y conversaciones y muestra todas las métricas "
         "del negocio en un panel en tiempo real.")
    para(doc,
         "Este manual está dividido en dos partes: la instalación y configuración del sistema "
         "(dirigida a la persona que administra el servidor) y la operación cotidiana de la "
         "plataforma (dirigida al superadministrador y a los dueños de cada negocio).")

    h1(doc, "2. Roles de usuario")
    bullets(doc, [
        ("Superadministrador",
         "es el titular de la plataforma. Administra los negocios, los planes, los cobros y el "
         "soporte. No opera un negocio propio."),
        ("Dueño del negocio",
         "es el cliente que contrata el servicio. Configura su agente, conecta su WhatsApp y "
         "consulta sus métricas, pero solo ve la información de su propio negocio."),
    ])

    h1(doc, "3. Requisitos")
    bullets(doc, [
        "Un servidor virtual (VPS) con sistema operativo Linux y Docker instalado.",
        "Un nombre de dominio propio (por ejemplo, con Porkbun) apuntando al servidor.",
        "Una clave de la interfaz de programación de Anthropic (ANTHROPIC_API_KEY) para activar la "
        "inteligencia artificial.",
        "Una conexión de WhatsApp Business API por cada negocio que se conecte, ya sea mediante la "
        "nube de Meta o mediante un proveedor de mensajería como Twilio.",
        "De forma opcional: cuentas de Twilio y Retell (voz), de Instagram y de un proveedor de "
        "almacenamiento para las copias de seguridad.",
    ])

    h1(doc, "4. Instalación y despliegue del sistema")
    para(doc, "Pasos para poner la plataforma en funcionamiento en un servidor:")
    numbered(doc, [
        "Adquirir un dominio y un servidor virtual, e instalar Docker y Docker Compose en el "
        "servidor.",
        "Apuntar el dominio (registros DNS de tipo A) a la dirección IP del servidor.",
        "Copiar el código fuente de la plataforma al servidor.",
        "Crear el archivo de configuración «.env» con el dominio, el correo para los certificados, "
        "las contraseñas de la base de datos y las claves de los servicios.",
        "Copiar el archivo «backend/.env.production» como «backend/.env» y completar los valores "
        "indicados.",
        "Ejecutar el comando de despliegue: «docker compose -f docker-compose.deploy.yml up -d "
        "--build».",
        "Esperar a que el sistema genere automáticamente el certificado de seguridad (HTTPS) y "
        "verificar que el dominio abre la página de la plataforma.",
    ])
    para(doc,
         "Tras el despliegue, el sistema crea automáticamente la cuenta de superadministrador con "
         "las credenciales definidas en la configuración.")

    h1(doc, "5. Configuración inicial (claves y servicios)")
    para(doc,
         "Las claves de los servicios externos se colocan en el archivo de configuración del "
         "servidor. La clave imprescindible es la de Anthropic (ANTHROPIC_API_KEY); sin ella, el "
         "agente de inteligencia artificial no genera respuestas. Las demás integraciones (voz, "
         "Instagram, almacenamiento) son opcionales y el sistema funciona sin ellas, degradando "
         "esas funciones de forma controlada.")

    h1(doc, "6. Operación como superadministrador")
    para(doc,
         "El superadministrador inicia sesión y accede a un panel dividido en pestañas. Desde allí "
         "puede:")
    bullets(doc, [
        "Ver el tablero general con métricas e indicadores de la plataforma y el estado de las "
        "claves de los servicios.",
        "Crear un negocio nuevo, definiendo el dueño, el plan y una contraseña inicial.",
        "Buscar negocios, activarlos o desactivarlos, cambiar su plan, restablecer la contraseña "
        "del dueño, exportar sus datos o eliminarlos.",
        "Consultar el uso y el consumo de cada negocio (contactos, mensajes, llamadas, costo real "
        "e ingreso).",
        "Revisar los cobros pendientes, confirmar un pago recibido por Yape o transferencia y "
        "suspender a los morosos.",
        "Atender las solicitudes de recuperación de contraseña de los dueños de negocio.",
    ])

    h1(doc, "7. Operación como dueño de negocio")
    h2(doc, "7.1. Registro y prueba gratuita")
    para(doc,
         "El dueño se registra desde la página pública, indicando su nombre, el nombre y el tipo de "
         "su negocio, su correo y una contraseña. Obtiene automáticamente un periodo de prueba "
         "gratuito de catorce días, sin necesidad de tarjeta.")
    h2(doc, "7.2. Configuración guiada (onboarding)")
    para(doc,
         "Tras el registro, un asistente paso a paso le da la bienvenida, le ayuda a conectar su "
         "WhatsApp Business y le orienta para configurar su agente.")
    h2(doc, "7.3. Conexión de WhatsApp")
    numbered(doc, [
        "Crear una aplicación en la plataforma para desarrolladores de Meta y agregar el producto "
        "WhatsApp.",
        "Copiar el identificador del número de teléfono («Phone Number ID») y el token de acceso.",
        "Pegar esos datos en la sección de conexión de WhatsApp de la plataforma.",
        "Registrar la dirección del webhook y el token de verificación que muestra la plataforma en "
        "la configuración de Meta.",
    ])
    para(doc,
         "De forma alternativa, el negocio puede conectar su WhatsApp mediante un proveedor de "
         "mensajería como Twilio, que actúa como intermediario autorizado ante Meta. En ese caso se "
         "registra el número del negocio en Twilio y se configura la dirección del webhook de la "
         "plataforma; una vez aprobado el número, los clientes escriben directamente al WhatsApp del "
         "negocio y el agente responde de forma automática.")
    h2(doc, "7.4. Configuración del agente de inteligencia artificial")
    para(doc,
         "En la sección «Agente IA», el dueño define el nombre del asistente, el horario de "
         "atención, los servicios que ofrece y las preguntas frecuentes. Puede usar el botón "
         "«Probar agente» para conversar con él antes de activarlo.")
    h2(doc, "7.5. Uso del panel")
    bullets(doc, [
        "Dashboard: métricas del día, volumen de mensajes y embudo de clientes potenciales.",
        "Conversaciones: historial de los chats atendidos por el agente.",
        "Llamadas: registro de las llamadas de voz.",
        "Contactos: lista de clientes con su etapa dentro del embudo.",
        "Instagram: publicaciones generadas por la inteligencia artificial, listas para aprobar.",
        "Automatizaciones: seguimientos y recordatorios programados.",
        "Configuración: datos del negocio, apariencia (tema claro u oscuro y color de marca) y "
        "preferencias de la cuenta.",
    ])

    h1(doc, "8. Personalización de la apariencia")
    para(doc,
         "Cada cuenta puede elegir entre el modo claro y el oscuro, y seleccionar el color principal "
         "de su panel. Estas preferencias se guardan de forma independiente por cuenta, de modo que "
         "no se mezclan entre los distintos usuarios de la plataforma.")

    h1(doc, "9. Facturación y cobro")
    para(doc,
         "El cobro es mensual y por adelantado. Tras los catorce días de prueba, el dueño realiza el "
         "pago por Yape o transferencia y el superadministrador lo confirma desde su panel, lo que "
         "renueva el servicio por un mes. Si no se recibe el pago, el servicio se suspende "
         "automáticamente hasta regularizar la situación. Los planes vigentes son: Inicial (S/ 149 "
         "mensuales), Básico (S/ 249 mensuales), Profesional (S/ 449 mensuales) y Empresarial "
         "(S/ 799 mensuales).")

    h1(doc, "10. Cierre de sesión y seguridad")
    para(doc,
         "Cada usuario inicia sesión con su correo y contraseña. Si un dueño olvida su contraseña, "
         "puede solicitar su recuperación desde la página de inicio de sesión; el "
         "superadministrador verifica la identidad y genera una nueva contraseña. Toda la "
         "información viaja cifrada mediante HTTPS y los datos de cada negocio están aislados de los "
         "demás.")

    h1(doc, "11. Soporte")
    para(doc,
         f"Para consultas sobre la plataforma, el contacto es {AUTOR}, a través del correo "
         f"{CORREO}.")

    h1(doc, "12. Créditos")
    para(doc,
         f"AgentePro fue desarrollado en colaboración por estudiantes de la {UNIVERSIDAD} —"
         f"{AUTOR} (autor principal y titular), {COAUTOR2} y {COAUTOR3} (coautores)— bajo la asesoría "
         f"del docente {ASESOR}.")

    doc.save(path)
    print("OK:", path)


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    memoria_descriptiva(os.path.join(here, "Memoria_Descriptiva_AgentePro.docx"))
    manual_usuario(os.path.join(here, "Manual_de_Usuario_AgentePro.docx"))
