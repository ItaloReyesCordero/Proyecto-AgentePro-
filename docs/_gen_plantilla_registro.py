# -*- coding: utf-8 -*-
"""Rellena la 'Plantilla de Registro General de Software' con el proyecto real
AgentePro: Anexo 1 (ejemplar con codigo fuente representativo real) y Anexo 2
(descripcion tecnica completa). Genera un documento nuevo y completo.
Ejecutar: python "_gen_plantilla_registro.py"
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

AZUL = RGBColor(0x1D, 0x4E, 0xD8)
GRIS = RGBColor(0x55, 0x55, 0x55)
NEGRO = RGBColor(0x11, 0x11, 0x11)
CODE_BG = "F2F4F8"

AUTOR = "Italo Eduardo Reyes Cordero"
DNI = "75220834"
COAUTOR2 = "Jack Joshua Bendezu Lagos"
DNI2 = "73940475"
COAUTOR3 = "Dickmar Wilber Julca Laureano"
DNI3 = "73086197"
ASESOR = "Maglioni Arana Caparachin"
UNIVERSIDAD = "Universidad Continental"
CIUDAD = "Huancayo, Perú"
FECHA = "Huancayo, 5 de junio de 2026"

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))

# (ruta relativa al proyecto, etiqueta a mostrar, limite de lineas)
ARCHIVOS = [
    ("agentepro/backend/app/main.py",
     "backend/app/main.py — Arranque de la aplicación y servidor del frontend", 92),
    ("agentepro/backend/app/core/tenant_scope.py",
     "backend/app/core/tenant_scope.py — Aislamiento de datos entre empresas (multiempresa)", 76),
    ("agentepro/backend/app/core/security.py",
     "backend/app/core/security.py — Seguridad: contraseñas y tokens de sesión", 72),
    ("agentepro/backend/app/utils/encryption.py",
     "backend/app/utils/encryption.py — Cifrado de credenciales sensibles", 40),
    ("agentepro/backend/app/utils/helpers.py",
     "backend/app/utils/helpers.py — Utilidades de dominio (zona horaria, etapa del lead)", 72),
    ("agentepro/backend/app/models/tenant.py",
     "backend/app/models/tenant.py — Modelo de datos del negocio (tenant)", 84),
    ("agentepro/backend/app/services/ai/agent.py",
     "backend/app/services/ai/agent.py — Motor del agente de inteligencia artificial", 96),
    ("agentepro/backend/app/services/whatsapp/webhook_handler.py",
     "backend/app/services/whatsapp/webhook_handler.py — Flujo de atención automática", 80),
    ("agentepro/frontend/src/lib/api.ts",
     "frontend/src/lib/api.ts — Cliente HTTP de la interfaz", 57),
    ("agentepro/frontend/src/stores/theme.store.ts",
     "frontend/src/stores/theme.store.ts — Tema y color de marca por cuenta", 60),
]


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.font.color.rgb = NEGRO
    rpr = st.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:ascii"), "Calibri")
    rf.set(qn("w:hAnsi"), "Calibri")


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text); r.font.color.rgb = AZUL; r.font.size = Pt(15); r.bold = True


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); r.font.color.rgb = NEGRO; r.font.size = Pt(12.5); r.bold = True


def para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullets(doc, items):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(it[0] + ": ").bold = True
            p.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")


def kv_table(doc, rows, headers=("Campo", "Detalle")):
    t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = t.rows[0].cells
    h[0].paragraphs[0].add_run(headers[0]).bold = True
    h[1].paragraphs[0].add_run(headers[1]).bold = True
    for a, b in rows:
        c = t.add_row().cells
        c[0].paragraphs[0].add_run(str(a)).bold = True
        c[1].paragraphs[0].add_run(str(b))
    doc.add_paragraph()
    return t


def shade(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def code_block(doc, lines):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(0); pf.space_before = Pt(0); pf.line_spacing = 1.0
    shade(p, CODE_BG)
    n = len(lines)
    for i, line in enumerate(lines):
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"; r.font.size = Pt(8.5)
        rpr = r._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:ascii"), "Consolas"); rf.set(qn("w:hAnsi"), "Consolas")
        if i < n - 1:
            r.add_break()
    doc.add_paragraph()
    return p


def read_lines(relpath, limit):
    path = os.path.join(ROOT, relpath)
    with open(path, encoding="utf-8") as f:
        raw = f.read().splitlines()
    out = [ln.rstrip() for ln in raw[:limit]]
    truncated = len(raw) > limit
    return out, truncated, len(raw)


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AGENTEPRO"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = AZUL
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Plantilla de Registro General de Software"); r.bold = True; r.font.size = Pt(18)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Ejemplar del software y descripción técnica")
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GRIS
    for _ in range(4):
        doc.add_paragraph()
    for k, v in [
        ("Título de la obra", "AgentePro"),
        ("Tipo de obra", "Programa de ordenador (plataforma web SaaS)"),
        ("Autor principal y titular", f"{AUTOR} — DNI N.º {DNI}"),
        ("Coautores", f"{COAUTOR2} (DNI N.º {DNI2}); {COAUTOR3} (DNI N.º {DNI3})"),
        ("Docente asesor", ASESOR),
        ("Institución", UNIVERSIDAD),
        ("Lugar y fecha", FECHA),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{k}: ").bold = True
        p.add_run(v)
    doc.add_page_break()


def build(path):
    doc = Document(); base_style(doc)
    cover(doc)

    # ───────────────────────── ANEXO 1 ─────────────────────────
    h1(doc, "ANEXO 1: Ejemplar del software")
    kv_table(doc, [
        ("Título de la obra", "AgentePro"),
        ("Tipo de software", "Plataforma web SaaS multiempresa de automatización de atención al "
                              "cliente con inteligencia artificial"),
        ("Lenguajes de programación", "Python, TypeScript, JavaScript y SQL"),
        ("Entorno / Framework", "Backend: FastAPI, SQLAlchemy, Socket.IO, Celery · Frontend: React, "
                                "Vite, Tailwind CSS · Datos: PostgreSQL (Alembic) · Infraestructura: "
                                "Docker, Caddy, Redis"),
        ("Autores", f"{AUTOR} (principal); {COAUTOR2}; {COAUTOR3}"),
        ("Titular", f"{AUTOR} — DNI N.º {DNI}"),
    ])
    para(doc,
         "A continuación se incluye una muestra representativa del código fuente del software, tomada "
         "directamente del proyecto. Se han seleccionado los módulos que reflejan el aporte propio y "
         "original de los autores —entre ellos, el aislamiento de datos entre empresas, el motor del "
         "agente de inteligencia artificial, la seguridad y el flujo de atención automática—. La "
         "muestra no contiene contraseñas, claves ni datos sensibles.")

    for rel, etiqueta, limit in ARCHIVOS:
        lines, truncated, total = read_lines(rel, limit)
        h2(doc, etiqueta)
        if truncated:
            para(doc, f"Fragmento representativo (se muestran las primeras {limit} líneas de {total}).",
                 justify=False)
        code_block(doc, lines)

    para(doc,
         "El código fuente completo y representativo se entrega, además, en el archivo comprimido "
         "«Codigo_Fuente_AgentePro.zip», que forma parte del expediente.")
    doc.add_page_break()

    # ───────────────────────── ANEXO 2 ─────────────────────────
    h1(doc, "ANEXO 2: Descripción técnica o resumen funcional")
    kv_table(doc, [
        ("Título del software", "AgentePro"),
        ("Versión", "1.0"),
        ("Autores", f"{AUTOR} (autor principal y titular); {COAUTOR2}; {COAUTOR3}"),
        ("Año de creación", "2026"),
        ("Lenguajes / Tecnologías empleadas",
         "Python (FastAPI, SQLAlchemy, Socket.IO, Celery), TypeScript y JavaScript (React, Vite, "
         "Tailwind CSS), SQL (PostgreSQL, Alembic), Docker, Caddy y Redis"),
        ("Tipo de obra", "Plataforma web (software)"),
    ])

    h2(doc, "1. Descripción general")
    para(doc,
         "AgentePro es una plataforma web de software como servicio (SaaS), multiempresa, que permite "
         "a un negocio automatizar la atención a sus clientes y su gestión comercial mediante "
         "inteligencia artificial. El sistema atiende de forma automática y permanente los mensajes de "
         "WhatsApp e Instagram, contesta y realiza llamadas de voz, registra contactos y "
         "conversaciones, califica clientes potenciales, agenda citas, genera contenido para redes "
         "sociales y muestra todas las métricas del negocio en un panel de control en tiempo real. "
         "Cada negocio cliente opera de forma aislada y segura dentro de la misma plataforma.")

    h2(doc, "2. Estructura y funcionamiento")
    para(doc, "El software se organiza en una arquitectura cliente-servidor con las siguientes capas:")
    bullets(doc, [
        ("Capa de presentación (frontend)", "interfaz web desarrollada en React y TypeScript "
         "(empaquetada con Vite y estilizada con Tailwind CSS); comprende la página pública, el panel "
         "del negocio y el panel de superadministración."),
        ("Capa lógica (backend)", "servidor en Python con FastAPI; implementa la lógica de negocio, la "
         "seguridad, el aislamiento de datos entre empresas, la comunicación en tiempo real "
         "(Socket.IO), las tareas en segundo plano (Celery) y la integración con servicios externos "
         "mediante una interfaz de programación (API) y webhooks."),
        ("Capa de datos", "base de datos PostgreSQL, con un esquema versionado mediante migraciones "
         "(Alembic)."),
        ("Infraestructura", "contenedores Docker, servidor web Caddy con HTTPS automático, memoria "
         "caché Redis y un servicio de copias de seguridad automáticas."),
    ])
    para(doc,
         "El flujo principal funciona así: un cliente escribe al WhatsApp o Instagram del negocio; la "
         "plataforma recibe el evento mediante un webhook y lo asocia al negocio correspondiente; el "
         "motor de inteligencia artificial genera la respuesta según la configuración del agente; la "
         "respuesta se envía al cliente, la conversación se registra y el sistema califica al cliente, "
         "agenda citas y programa seguimientos. Todo queda visible en el panel en tiempo real.")

    h2(doc, "3. Características técnicas")
    bullets(doc, [
        "Arquitectura multiempresa con aislamiento de datos forzado a nivel de aplicación: cada "
        "negocio solo accede a su propia información.",
        "Autenticación con tokens de sesión (JWT) y control de roles (superadministrador y dueño de "
        "negocio).",
        "Agente conversacional con inteligencia artificial (modelos de lenguaje de Anthropic) para "
        "WhatsApp, Instagram y voz.",
        "Integración con WhatsApp Business API (nube de Meta o proveedor Twilio), telefonía de voz "
        "(Twilio y Retell) e Instagram.",
        "Recepción y verificación de eventos externos mediante webhooks con validación de firma.",
        "Comunicación en tiempo real con Socket.IO y tareas programadas en segundo plano con Celery.",
        "Cifrado de credenciales sensibles y transmisión segura mediante HTTPS.",
        "Panel de superadministración con gestión de negocios, planes, uso, costos y cobros.",
        "Cobro mensual por adelantado, con periodo de prueba gratuito y suspensión automática por "
        "falta de pago.",
        "Batería de pruebas automatizadas (unitarias, de integración, funcionales y de seguridad).",
    ])

    h2(doc, "4. Originalidad")
    para(doc,
         "El elemento original y distintivo de la obra reside en la combinación de su arquitectura "
         "multiempresa con aislamiento de datos forzado a nivel de aplicación, el aprovisionamiento "
         "automático de los canales de comunicación de cada negocio (mensajería, voz y redes "
         "sociales) y la orquestación de un motor de inteligencia artificial que responde, califica y "
         "da seguimiento a los clientes de cada negocio de forma autónoma. La selección, disposición y "
         "organización de estos componentes, así como el código fuente que los implementa, "
         "constituyen una creación intelectual propia de los autores.")

    h2(doc, "5. Entorno de ejecución")
    bullets(doc, [
        ("Sistema operativo del servidor", "Linux (con Docker)."),
        ("Cliente", "navegador web moderno (Google Chrome, Mozilla Firefox), en computadora o "
         "dispositivo móvil."),
        ("Base de datos", "PostgreSQL."),
        ("Framework / motor", "FastAPI (servidor) y React (interfaz)."),
        ("Servicios externos", "Anthropic (inteligencia artificial), Meta o Twilio (WhatsApp), "
         "Twilio y Retell (voz), opcionalmente Instagram y almacenamiento para copias de seguridad."),
        ("Requisitos", "dominio propio con HTTPS y la clave de Anthropic para activar la inteligencia "
         "artificial."),
    ])

    h2(doc, "6. Archivos adjuntos")
    bullets(doc, [
        "Capturas de pantalla de la página pública, el inicio de sesión, el panel del negocio y el "
        "panel de superadministración.",
        "Diagramas de arquitectura y del flujo de atención automática.",
        "Código fuente representativo (Codigo_Fuente_AgentePro.zip).",
    ])
    para(doc,
         f"Documento elaborado por el equipo de desarrollo de AgentePro, estudiantes de la "
         f"{UNIVERSIDAD}, bajo la asesoría del docente {ASESOR}. {FECHA}.")

    doc.save(path)
    print("OK:", path)


if __name__ == "__main__":
    build(os.path.join(HERE, "Plantilla Registro General de Software - AgentePro (rellenado).docx"))
