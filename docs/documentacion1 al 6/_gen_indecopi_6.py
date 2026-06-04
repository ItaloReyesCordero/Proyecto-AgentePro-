# -*- coding: utf-8 -*-
"""Genera los 6 documentos de registro ante INDECOPI, YA RELLENADOS con los
datos reales de AgentePro (autor único, persona natural). Crea versiones
finales junto a las plantillas originales.
Ejecutar: python "_gen_indecopi_6.py"
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

AZUL = RGBColor(0x1D, 0x4E, 0xD8)
GRIS = RGBColor(0x55, 0x55, 0x55)
NEGRO = RGBColor(0x11, 0x11, 0x11)

AUTOR = "Italo Eduardo Reyes Cordero"
CORREO = "italoreyescordero1@gmail.com"
DNI = "75220834"
DOMICILIO = "Jr. Grau N.º 419, distrito de Jauja, provincia de Jauja, departamento de Junín"
TELEFONO = "916085873"
CIUDAD = "Huancayo, Perú"
CIUDAD_FIRMA = "Huancayo"
ANIO = "2026"
SOFTWARE = "AgentePro"


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.font.color.rgb = NEGRO
    rpr = st.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:ascii"), "Calibri")
    rf.set(qn("w:hAnsi"), "Calibri")


def doc_header(doc, numero, titulo):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"DOCUMENTO {numero}"); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = GRIS
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo.upper()); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = AZUL
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Software: {SOFTWARE}  ·  Autor y titular: {AUTOR}  ·  {CIUDAD}, {ANIO}")
    r.font.size = Pt(10); r.font.color.rgb = GRIS
    doc.add_paragraph()


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text); r.font.color.rgb = AZUL; r.font.size = Pt(14); r.bold = True


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); r.font.color.rgb = NEGRO; r.font.size = Pt(12); r.bold = True


def para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullets(doc, items):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(it[0] + ": ").bold = True
            p.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def kv_table(doc, rows, headers=("Campo", "Detalle")):
    t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = t.rows[0].cells
    h[0].paragraphs[0].add_run(headers[0]).bold = True
    h[1].paragraphs[0].add_run(headers[1]).bold = True
    for a, b in rows:
        c = t.add_row().cells
        c[0].paragraphs[0].add_run(str(a)).bold = True
        c[1].paragraphs[0].add_run(str(b))
    doc.add_paragraph()
    return t


def firma(doc, cargo):
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_______________________________").bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(AUTOR + "\n").bold = True
    p.add_run(cargo + "\n")
    p.add_run(f"DNI N.º {DNI}   ·   " + f"{CIUDAD_FIRMA}, ____ de __________ de {ANIO}")


# ════════════════════ DOCUMENTO 1 — FICHA TÉCNICA ════════════════════
def doc1(path):
    d = Document(); base_style(d)
    doc_header(d, 1, "Ficha Técnica del Software")
    para(d, "La presente ficha técnica identifica de manera clara y técnica el programa de ordenador "
            "que se registra ante INDECOPI, describiendo su denominación, lenguajes, funcionalidad, "
            "fecha de creación y arquitectura.")

    h1(d, "1. Título del software")
    para(d, "El presente programa de ordenador se denomina «AgentePro». La denominación identifica de "
            "forma única a la plataforma y la distingue de otros sistemas, evitando nombres genéricos. "
            "AgentePro es el nombre comercial y técnico con el que se reconoce la obra.")

    h1(d, "2. Lenguajes de programación utilizados")
    para(d, "El software ha sido desarrollado principalmente en Python (para el servidor y la lógica "
            "de negocio), TypeScript y JavaScript con la biblioteca React (para la interfaz de "
            "usuario) y SQL sobre PostgreSQL (para la gestión de la base de datos). Se emplearon "
            "tecnologías complementarias como FastAPI y Socket.IO en el servidor, Vite y Tailwind CSS "
            "en la interfaz, y Docker para el empaquetado y despliegue. La elección de estos lenguajes "
            "responde a criterios de escalabilidad, seguridad, rendimiento y facilidad de "
            "mantenimiento.")
    kv_table(d, [
        ("Servidor (backend)", "Python (FastAPI, SQLAlchemy, Socket.IO, Celery)"),
        ("Interfaz (frontend)", "TypeScript y JavaScript (React, Vite, Tailwind CSS)"),
        ("Base de datos", "SQL sobre PostgreSQL (migraciones con Alembic)"),
        ("Infraestructura", "Docker, Docker Compose, servidor web Caddy, Redis"),
    ], headers=("Capa", "Lenguajes y tecnologías"))

    h1(d, "3. Funcionalidad principal")
    para(d, "AgentePro tiene como funcionalidad principal la automatización de la atención al cliente "
            "y la gestión comercial de un negocio mediante inteligencia artificial. El sistema atiende "
            "de forma automática y permanente los mensajes de WhatsApp e Instagram, contesta y realiza "
            "llamadas de voz, registra contactos y conversaciones, califica clientes potenciales, "
            "agenda citas, genera contenido para redes sociales, ejecuta seguimientos automáticos y "
            "muestra todas las métricas del negocio en un panel de control en tiempo real. Es una "
            "plataforma multiempresa: cada negocio cliente opera de forma aislada y segura.")

    h1(d, "4. Fecha de creación")
    para(d, "La primera versión funcional del software fue concluida en mayo de 2026, tras un proceso "
            "de desarrollo realizado durante el primer semestre de 2026, que incluyó las fases de "
            "análisis, diseño, programación, pruebas y documentación.")

    h1(d, "5. Arquitectura del sistema")
    para(d, "AgentePro se estructura bajo una arquitectura cliente-servidor. El servidor, desarrollado "
            "en Python con FastAPI, implementa la lógica de negocio, la seguridad, el aislamiento de "
            "datos entre empresas, la comunicación en tiempo real y la integración con servicios "
            "externos. La interfaz, desarrollada en React, ofrece el panel del negocio y el panel de "
            "superadministración. La base de datos PostgreSQL almacena la información de forma "
            "estructurada. Todo el sistema se despliega en contenedores Docker, con un servidor web "
            "Caddy que provee cifrado HTTPS automático. La arquitectura es modular y está preparada "
            "para funcionar en la nube.")

    h1(d, "6. Observaciones finales")
    para(d, "La presente ficha técnica constituye un resumen del software AgentePro y forma parte del "
            "paquete de documentos para su registro como programa de ordenador ante INDECOPI, conforme "
            "al Decreto Legislativo N.º 822.")
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 2 — EJEMPLAR DEL SOFTWARE ════════════════════
def doc2(path):
    d = Document(); base_style(d)
    doc_header(d, 2, "Ejemplar del Software")
    para(d, "El ejemplar del software constituye la evidencia técnica de la obra. Demuestra su "
            "existencia, funcionalidad y originalidad mediante el código fuente, un manual técnico y "
            "material gráfico de las interfaces.")

    h1(d, "1. Código fuente")
    para(d, "Se adjunta un archivo comprimido denominado «Codigo_Fuente_AgentePro.zip», que contiene "
            "una versión representativa del código fuente del software, organizada de la siguiente "
            "manera:")
    bullets(d, [
        ("Carpeta backend", "código del servidor en Python (FastAPI): modelos de datos, servicios, "
         "puntos de acceso de la interfaz de programación, webhooks, tareas en segundo plano y "
         "migraciones de la base de datos."),
        ("Carpeta frontend", "código de la interfaz en TypeScript y React: páginas, componentes, "
         "estados y utilidades."),
        ("Archivos de configuración y despliegue", "definiciones de contenedores Docker y del "
         "servidor web."),
        ("Documentación técnica", "archivos de apoyo en la carpeta de documentación."),
    ])
    para(d, "El archivo comprimido representa fielmente la autoría del titular y está limpio de "
            "credenciales, contraseñas y datos sensibles, por lo que constituye un ejemplar suficiente "
            "para acreditar la existencia del software ante INDECOPI.")

    h1(d, "2. Manual técnico")
    para(d, "El manual técnico guía la instalación, configuración y uso del software. Se adjunta de "
            "forma completa en el documento «Manual de Usuario de AgentePro». A continuación se "
            "resumen sus apartados principales:")
    h2(d, "2.1. Requisitos del sistema")
    bullets(d, [
        "Servidor con sistema operativo Linux y Docker instalado.",
        "Nombre de dominio propio apuntando al servidor.",
        "Clave de la interfaz de programación de Anthropic para activar la inteligencia artificial.",
        "Cuenta de WhatsApp Business API por cada negocio, conectada mediante la nube de Meta o "
        "mediante un proveedor de mensajería como Twilio.",
        "Navegador web moderno (Google Chrome o Mozilla Firefox).",
    ])
    h2(d, "2.2. Instalación y despliegue")
    numbered(d, [
        "Instalar Docker y Docker Compose en el servidor.",
        "Apuntar el dominio (registro DNS de tipo A) a la dirección IP del servidor.",
        "Copiar el código fuente al servidor y crear el archivo de configuración «.env».",
        "Ejecutar el comando de despliegue: docker compose -f docker-compose.deploy.yml up -d --build.",
        "Esperar a que el sistema genere el certificado de seguridad (HTTPS) y verificar el dominio.",
    ])
    h2(d, "2.3. Uso básico")
    bullets(d, [
        "Iniciar sesión con correo y contraseña.",
        "Conectar el WhatsApp del negocio y configurar el agente de inteligencia artificial.",
        "Consultar conversaciones, llamadas, contactos y métricas en el panel.",
    ])
    h2(d, "2.4. Mantenimiento")
    bullets(d, [
        "Las copias de seguridad de la base de datos se realizan de forma automática y periódica.",
        "Las actualizaciones se aplican reconstruyendo los contenedores.",
    ])

    h1(d, "3. Capturas de pantalla")
    para(d, "Se adjuntan capturas de pantalla que muestran las principales interfaces del software:")
    bullets(d, [
        "Página pública de presentación, con la marca y los planes.",
        "Pantalla de inicio de sesión.",
        "Panel del negocio (dashboard): métricas, volumen de mensajes y embudo de clientes.",
        "Sección de conversaciones atendidas por el agente.",
        "Panel de superadministración: gestión de negocios, uso y cobros.",
    ])
    para(d, "Las capturas permiten visualizar la funcionalidad principal del sistema y constituyen "
            "evidencia gráfica de su existencia. Las imágenes no contienen datos personales reales.")

    h1(d, "4. Observaciones finales")
    para(d, "El ejemplar del software AgentePro constituye evidencia suficiente para acreditar su "
            "existencia y funcionalidad ante INDECOPI. Junto con la ficha técnica, la declaración "
            "jurada de autoría y los anexos técnicos, conforma el paquete completo de registro.")
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 3 — DECLARACIÓN JURADA DE AUTORÍA ════════════════════
def doc3(path):
    d = Document(); base_style(d)
    doc_header(d, 3, "Declaración Jurada de Autoría")
    para(d, "La presente Declaración Jurada de Autoría reconoce la autoría del software, declara su "
            "originalidad y deja constancia de la titularidad de los derechos.")

    h1(d, "1. Identificación del declarante")
    para(d, f"Yo, {AUTOR}, identificado con Documento Nacional de Identidad (DNI) N.º {DNI}, con "
            f"domicilio en {DOMICILIO}, de nacionalidad peruana, en calidad de autor y titular del "
            f"software denominado «{SOFTWARE}», declaro bajo juramento lo siguiente:")

    h1(d, "2. Reconocimiento de autoría")
    para(d, "Declaro ser el autor del software mencionado, desarrollado en los lenguajes de "
            "programación Python, TypeScript, JavaScript y SQL, cuya funcionalidad principal consiste "
            "en automatizar la atención al cliente y la gestión comercial de un negocio mediante "
            "inteligencia artificial. El software fue concebido, diseñado y programado de manera "
            "individual por el suscrito, quien desempeñó de forma integral los roles de análisis, "
            "diseño de arquitectura, programación del servidor (backend), programación de la interfaz "
            "(frontend), diseño de la base de datos y pruebas de calidad.")

    h1(d, "3. Autoría individual")
    para(d, "Declaro que el software es una obra de autoría individual. No existen coautores; la "
            "totalidad del código fuente y del diseño de la obra es producto de mi trabajo creativo "
            "personal.")

    h1(d, "4. Declaración de originalidad")
    para(d, "Declaro que la obra es original, producto de mi trabajo creativo, y que no infringe "
            "derechos de terceros. El software no ha sido copiado ni adaptado de programas ajenos sin "
            "autorización, y constituye una creación independiente. El uso de bibliotecas y marcos de "
            "trabajo de terceros se realiza conforme a sus respectivas licencias de código abierto, "
            "sin que ello afecte la originalidad de la obra desarrollada.")

    h1(d, "5. Titularidad de los derechos")
    para(d, "En mi condición de autor y persona natural, me corresponden tanto los derechos morales "
            "como los derechos patrimoniales del software, de conformidad con el Decreto Legislativo "
            "N.º 822, Ley sobre el Derecho de Autor. No existe empresa ni tercero titular de los "
            "derechos patrimoniales. Solicito y autorizo el registro de la obra a mi nombre ante "
            "INDECOPI.")

    h1(d, "6. Cláusula de juramento")
    para(d, "Declaro bajo juramento la veracidad de lo expuesto en el presente documento, "
            "comprometiéndome a responder legalmente en caso de falsedad.")

    h1(d, "7. Firma y fecha")
    firma(d, "Autor y titular de la obra")
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 4 — LISTA DE AUTORES Y ROLES ════════════════════
def doc4(path):
    d = Document(); base_style(d)
    doc_header(d, 4, "Lista de Autores y Roles")
    para(d, "El presente documento identifica al autor del software y detalla los roles que "
            "desempeñó durante su creación.")

    h1(d, "1. Identificación del software")
    para(d, f"El presente documento corresponde al software denominado «{SOFTWARE}», desarrollado de "
            f"manera individual por {AUTOR}. La lista que se presenta a continuación detalla al autor, "
            f"sus roles y el periodo de participación en el proyecto.")

    h1(d, "2. Tabla de autores y roles")
    t = d.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
    heads = ("Nombre completo / DNI", "Rol en el proyecto", "Componentes desarrollados", "Periodo")
    for i, h in enumerate(heads):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    row = t.add_row().cells
    row[0].paragraphs[0].add_run(f"{AUTOR}\nDNI N.º {DNI}")
    row[1].paragraphs[0].add_run("Autor único y desarrollador integral (full-stack): análisis, "
                                 "diseño de arquitectura, programación, base de datos y pruebas.")
    row[2].paragraphs[0].add_run("Servidor (backend en FastAPI), interfaz (frontend en React), "
                                 "esquema y migraciones de base de datos, integraciones, webhooks, "
                                 "panel de administración y batería de pruebas automatizadas.")
    row[3].paragraphs[0].add_run("Primer semestre de 2026")
    d.add_paragraph()

    h1(d, "3. Descripción de los aportes")
    para(d, "El autor diseñó la arquitectura multiempresa con aislamiento de datos, programó la "
            "totalidad de la lógica de negocio del servidor, implementó la interfaz de usuario y el "
            "panel de superadministración, modeló la base de datos y sus migraciones, integró los "
            "servicios de inteligencia artificial, mensajería y telefonía, y elaboró la batería de "
            "pruebas automatizadas que verifican el funcionamiento del sistema. Al tratarse de una "
            "obra de autoría individual, todos los roles fueron desempeñados por la misma persona.")

    h1(d, "4. Validación")
    para(d, f"El suscrito, {AUTOR}, en calidad de autor y titular de los derechos del software, valida "
            f"la presente lista de autores y roles. Este documento forma parte del paquete de registro "
            f"ante INDECOPI.")

    h1(d, "5. Firma y fecha")
    firma(d, "Autor y titular de la obra")
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 5 — TITULARIDAD (PERSONA NATURAL) ════════════════════
def doc5(path):
    d = Document(); base_style(d)
    doc_header(d, 5, "Titularidad de la Obra (Persona Natural)")
    para(d, "El presente documento acredita la titularidad del software. A diferencia del modelo de "
            "empresa titular, en este caso el titular de los derechos es una persona natural: el "
            "propio autor. Por ello, no aplica la figura de representación legal de una empresa.")

    h1(d, "1. Naturaleza de la titularidad")
    para(d, f"El software «{SOFTWARE}» es una obra de autoría individual cuya titularidad recae "
            f"íntegramente en una persona natural, {AUTOR}. No existe persona jurídica (empresa) "
            f"titular de los derechos ni representante legal que la gestione; el autor actúa por "
            f"derecho propio.")

    h1(d, "2. Identificación del titular")
    kv_table(d, [
        ("Titular y autor", AUTOR),
        ("Tipo de titular", "Persona natural"),
        ("Documento de identidad", f"DNI N.º {DNI}"),
        ("Domicilio", DOMICILIO),
        ("Nacionalidad", "Peruana"),
        ("Correo electrónico", CORREO),
        ("Teléfono", TELEFONO),
    ])

    h1(d, "3. Declaración de titularidad")
    para(d, f"{AUTOR} declara ser el titular de los derechos morales y patrimoniales del software "
            f"«{SOFTWARE}», desarrollado de manera individual. En ejercicio de sus derechos "
            f"patrimoniales, se reserva las facultades de reproducción, distribución, comunicación "
            f"pública y transformación de la obra, conforme al Decreto Legislativo N.º 822, Ley sobre "
            f"el Derecho de Autor.")

    h1(d, "4. Facultades para el registro")
    para(d, "En su condición de autor y titular, se encuentra plenamente facultado para realizar "
            "todos los actos necesarios para el registro de la obra ante INDECOPI, incluyendo la "
            "presentación de la solicitud, los anexos, el ejemplar del software y la declaración "
            "jurada de autoría, así como para suscribir, en el futuro, contratos de licencia o cesión "
            "de derechos sobre la obra.")

    h1(d, "5. Nota sobre una eventual empresa titular")
    para(d, "En caso de que, en el futuro, el autor constituya una empresa (por ejemplo, una "
            "E.I.R.L. o una S.A.C.) y desee que la titularidad de los derechos patrimoniales "
            "corresponda a dicha empresa, deberá formalizarse una cesión de derechos del autor a la "
            "empresa y actualizarse el registro. A la fecha del presente documento, la titularidad "
            "corresponde a la persona natural antes identificada.")

    h1(d, "6. Cláusula de responsabilidad")
    para(d, "El titular asume plena responsabilidad por la veracidad de la información contenida en "
            "el presente documento y por la autenticidad de los anexos que se adjuntan, "
            "comprometiéndose a responder conforme a la legislación vigente en caso de falsedad.")

    h1(d, "7. Firma y fecha")
    firma(d, "Autor y titular de la obra (persona natural)")
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 6 — ANEXOS TÉCNICOS ════════════════════
def doc6(path):
    d = Document(); base_style(d)
    doc_header(d, 6, "Anexos Técnicos")
    para(d, "Los anexos técnicos constituyen material complementario que respalda la descripción y el "
            "ejemplar del software, demostrando su estructura interna, su funcionamiento, sus "
            "interfaces y los procesos de calidad aplicados.")

    h1(d, "1. Flujo principal del sistema")
    para(d, "El flujo principal de AgentePro, para la atención automática de un cliente, comprende las "
            "siguientes etapas:")
    numbered(d, [
        "Un cliente envía un mensaje al WhatsApp o Instagram del negocio.",
        "La plataforma recibe el evento mediante un webhook y lo asocia al negocio correspondiente.",
        "El motor de inteligencia artificial genera una respuesta según la configuración del agente.",
        "La respuesta se envía al cliente y la conversación se registra en la base de datos.",
        "El sistema califica al cliente potencial, agenda citas y programa seguimientos según "
        "corresponda.",
        "El negocio visualiza la actividad y las métricas en su panel de control en tiempo real.",
    ])
    para(d, "Se recomienda adjuntar un diagrama de flujo (por ejemplo, en notación BPMN o UML) que "
            "represente gráficamente este proceso.")

    h1(d, "2. Arquitectura del software")
    para(d, "AgentePro se estructura bajo una arquitectura cliente-servidor, con los siguientes "
            "componentes:")
    bullets(d, [
        ("Servidor (backend)", "desarrollado en Python con FastAPI; gestiona la lógica de negocio, la "
         "seguridad, el aislamiento de datos entre empresas y las integraciones externas; expone una "
         "interfaz de programación y procesa eventos mediante webhooks; usa Socket.IO para la "
         "comunicación en tiempo real y Celery para las tareas en segundo plano."),
        ("Interfaz (frontend)", "desarrollada en TypeScript con React; comprende la página pública, el "
         "panel del negocio y el panel de superadministración."),
        ("Base de datos", "PostgreSQL, con un esquema versionado mediante migraciones."),
        ("Infraestructura", "contenedores Docker, servidor web Caddy con HTTPS automático, memoria "
         "caché Redis y un servicio de copias de seguridad automáticas."),
    ])
    para(d, "La arquitectura es modular y está preparada para funcionar en la nube. Se recomienda "
            "adjuntar un diagrama de arquitectura como anexo gráfico.")

    h1(d, "3. Capturas de pantalla de módulos clave")
    bullets(d, [
        "Página pública de presentación con la marca y los planes.",
        "Pantalla de inicio de sesión.",
        "Panel del negocio: métricas, volumen de mensajes y embudo de clientes potenciales.",
        "Sección de conversaciones y de configuración del agente de inteligencia artificial.",
        "Panel de superadministración: gestión de negocios, uso y consumo, y cobros.",
    ])

    h1(d, "4. Resumen de pruebas de calidad (QA)")
    para(d, "El software cuenta con una batería de pruebas automatizadas que verifica su correcto "
            "funcionamiento. El resumen de resultados es el siguiente:")
    kv_table(d, [
        ("Pruebas de servidor ejecutadas con éxito", "543"),
        ("Pruebas deseleccionadas (requieren claves de servicios reales)", "5"),
        ("Tipos de prueba", "Unitarias, de integración (HTTP), funcionales y de seguridad"),
        ("Pruebas de interfaz", "Ejecutadas correctamente con Vitest"),
        ("Construcción de producción del frontend", "Sin errores"),
        ("Defectos críticos pendientes", "Ninguno"),
    ], headers=("Indicador", "Resultado"))
    para(d, "El detalle completo se encuentra en el documento «Plan de Pruebas de Software de "
            "AgentePro».")

    h1(d, "5. Documentación complementaria")
    bullets(d, [
        "Manual de Usuario de AgentePro (instalación, configuración y operación).",
        "Memoria Descriptiva de la obra.",
        "Plan de Pruebas de Software.",
        "Ficha Técnica del Software (Documento 1).",
    ])

    h1(d, "6. Observaciones finales")
    para(d, "Los anexos técnicos presentados constituyen evidencia adicional que respalda la "
            "originalidad y funcionalidad del software AgentePro, y complementan el paquete de "
            "documentos para su registro ante INDECOPI.")
    d.save(path); print("OK:", path)


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    p = lambda n: os.path.join(here, n)
    doc1(p("Documento 1 - Ficha técnica de software - AgentePro.docx"))
    doc2(p("Documento 2 - Ejemplar del software - AgentePro.docx"))
    doc3(p("Documento 3 - Declaración jurada de autoría - AgentePro.docx"))
    doc4(p("Documento 4 - Lista de autores y roles - AgentePro.docx"))
    doc5(p("Documento 5 - Titularidad de la obra - AgentePro.docx"))
    doc6(p("Documento 6 - Anexos técnicos - AgentePro.docx"))
