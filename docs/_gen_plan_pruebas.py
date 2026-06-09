# -*- coding: utf-8 -*-
"""Genera el Plan de Pruebas de Software de AgentePro en .docx.
Ejecutar: python _gen_plan_pruebas.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

AZUL = RGBColor(0x1D, 0x4E, 0xD8)
GRIS = RGBColor(0x55, 0x55, 0x55)
NEGRO = RGBColor(0x11, 0x11, 0x11)

AUTOR = "Italo Eduardo Reyes Cordero"
CORREO = "italoreyescordero1@gmail.com"
DNI = "75220834"
COAUTOR2 = "Jack Joshua Bendezu Lagos"
DNI2 = "73940475"
COAUTOR3 = "Dickmar Wilber Julca Laureano"
DNI3 = "73086197"
ASESOR = "Maglioni Arana Caparachin"
UNIVERSIDAD = "Universidad Continental"
CIUDAD = "Huancayo, Perú"
FECHA_FIRMA = "Huancayo, 5 de junio de 2026"
ANIO = "2026"


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.font.color.rgb = NEGRO
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = AZUL
    r.font.size = Pt(15)
    r.bold = True


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = NEGRO
    r.font.size = Pt(12.5)
    r.bold = True


def para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullets(doc, items):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(it[0] + ": ")
            r.bold = True
            p.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")


def kv_table(doc, rows, headers=("Campo", "Detalle")):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    hdr[0].paragraphs[0].add_run(headers[0]).bold = True
    hdr[1].paragraphs[0].add_run(headers[1]).bold = True
    for a, b in rows:
        c = t.add_row().cells
        c[0].paragraphs[0].add_run(str(a)).bold = True
        c[1].paragraphs[0].add_run(str(b))
    doc.add_paragraph()
    return t


def cases_table(doc, rows):
    headers = ("ID", "Módulo", "Caso de prueba", "Resultado esperado", "Estado")
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].add_run(h).bold = True
    for row in rows:
        c = t.add_row().cells
        for i, val in enumerate(row):
            run = c[i].paragraphs[0].add_run(str(val))
            if i == 0:
                run.bold = True
            for rr in c[i].paragraphs[0].runs:
                rr.font.size = Pt(9.5)
    doc.add_paragraph()
    return t


def title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AGENTEPRO"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = AZUL
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Plataforma SaaS de automatización de atención al cliente con inteligencia artificial")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRIS
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PLAN DE PRUEBAS DE SOFTWARE"); r.bold = True; r.font.size = Pt(20)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento de aseguramiento de la calidad"); r.font.size = Pt(12); r.font.color.rgb = GRIS
    for _ in range(6):
        doc.add_paragraph()
    for k, v in [("Software", "AgentePro"), ("Autor principal", AUTOR),
                 ("Coautores", f"{COAUTOR2}; {COAUTOR3}"), ("Año", ANIO),
                 ("Lugar", CIUDAD), ("Correo", CORREO), ("Versión del documento", "1.0")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{k}: "); r.bold = True
        p.add_run(v)
    doc.add_page_break()


def build(path):
    doc = Document()
    base_style(doc)
    title_page(doc)

    h1(doc, "1. Introducción")
    para(doc,
         "El presente Plan de Pruebas de Software describe la estrategia, el alcance, los tipos de "
         "prueba, el entorno y los casos empleados para verificar la calidad del software AgentePro. "
         "Su objetivo es asegurar que el sistema cumple con los requisitos funcionales y no "
         "funcionales definidos, y que opera de manera correcta, segura y estable antes de su puesta "
         "en producción.")

    h1(doc, "2. Objetivos de las pruebas")
    bullets(doc, [
        "Verificar que cada funcionalidad del sistema cumple con su propósito.",
        "Comprobar la correcta integración entre el servidor, la interfaz y la base de datos.",
        "Validar la seguridad: autenticación, autorización y aislamiento de datos entre negocios.",
        "Confirmar que los flujos críticos (registro, inicio de sesión, mensajería, cobros) "
        "funcionan de extremo a extremo.",
        "Detectar y corregir defectos antes del despliegue.",
    ])

    h1(doc, "3. Alcance")
    para(doc,
         "Las pruebas cubren el componente de servidor (backend en Python/FastAPI), el componente de "
         "interfaz (frontend en React/TypeScript) y la capa de datos (PostgreSQL). Se incluyen pruebas "
         "unitarias, de integración, funcionales, de seguridad y de aceptación. Las integraciones con "
         "servicios externos (modelos de inteligencia artificial, mensajería, telefonía) se prueban "
         "mediante simulaciones controladas, ya que dependen de claves y cuentas de terceros.")

    h1(doc, "4. Identificación del software bajo prueba")
    kv_table(doc, [
        ("Nombre", "AgentePro"),
        ("Tipo", "Aplicación web SaaS multiempresa"),
        ("Backend", "Python con FastAPI, SQLAlchemy y Socket.IO"),
        ("Frontend", "TypeScript y React con Vite"),
        ("Base de datos", "PostgreSQL"),
        ("Autor principal y titular", AUTOR),
        ("Coautores", f"{COAUTOR2}; {COAUTOR3}"),
    ])

    h1(doc, "5. Estrategia y tipos de prueba")
    bullets(doc, [
        ("Pruebas unitarias",
         "verifican funciones aisladas (utilidades, seguridad, cifrado, límites de tasa, validación "
         "de esquemas y propiedades de los modelos)."),
        ("Pruebas de integración (HTTP)",
         "ejercitan los puntos de acceso reales de la interfaz de programación con una base de datos "
         "en memoria, simulando peticiones completas de los usuarios."),
        ("Pruebas funcionales",
         "validan los flujos de negocio de principio a fin: registro, inicio de sesión, conexión de "
         "WhatsApp, configuración del agente, mensajería y cobros."),
        ("Pruebas de seguridad",
         "comprueban la autenticación obligatoria, la separación entre rol de superadministrador y "
         "rol de negocio, el aislamiento de datos entre empresas y la validación de firmas de los "
         "webhooks."),
        ("Pruebas de la interfaz",
         "verifican la lógica del frontend (tema, manejo de errores de la interfaz de programación) "
         "con un marco de pruebas dedicado."),
        ("Pruebas de aceptación",
         "validan, de forma manual, que el producto satisface las expectativas del usuario final."),
    ])

    h1(doc, "6. Entorno y herramientas de prueba")
    kv_table(doc, [
        ("Pruebas de backend", "Pytest y Pytest-asyncio, con cliente HTTP asíncrono (httpx)."),
        ("Base de datos de prueba", "SQLite en memoria (aislada por prueba)."),
        ("Cobertura de código", "Pytest-cov."),
        ("Pruebas de frontend", "Vitest con entorno jsdom."),
        ("Verificación de construcción", "Compilador de TypeScript y empaquetador Vite."),
        ("Contenedores", "Docker y Docker Compose para el entorno integral."),
    ], headers=("Aspecto", "Herramienta"))

    h1(doc, "7. Criterios de entrada y de salida")
    h2(doc, "7.1. Criterios de entrada")
    bullets(doc, [
        "El código de la funcionalidad a probar está terminado e integrado.",
        "El entorno de pruebas está configurado y disponible.",
        "Los casos de prueba están definidos.",
    ])
    h2(doc, "7.2. Criterios de salida")
    bullets(doc, [
        "El 100 % de los casos de prueba planificados se han ejecutado.",
        "No existen defectos críticos ni de severidad alta sin resolver.",
        "La construcción del frontend se completa sin errores.",
    ])

    h1(doc, "8. Casos de prueba representativos")
    para(doc,
         "A continuación se listan casos de prueba representativos del sistema. El estado «Aprobado» "
         "indica que el resultado obtenido coincidió con el resultado esperado.")
    cases_table(doc, [
        ("CP-01", "Autenticación", "Registrar un negocio nuevo con datos válidos",
         "Se crea la cuenta en periodo de prueba y se devuelven los tokens de sesión", "Aprobado"),
        ("CP-02", "Autenticación", "Iniciar sesión con contraseña incorrecta",
         "Se rechaza con un mensaje claro en español y código 401", "Aprobado"),
        ("CP-03", "Autenticación", "Acceder a un recurso de negocio sin sesión",
         "El sistema responde 401 (no autorizado)", "Aprobado"),
        ("CP-04", "Autorización", "Acceder a un recurso de administración sin ser superadministrador",
         "El sistema responde 403 (prohibido)", "Aprobado"),
        ("CP-05", "Aislamiento", "Consultar datos estando autenticado como otro negocio",
         "No se devuelven datos de negocios ajenos (conjunto vacío)", "Aprobado"),
        ("CP-06", "Agente IA", "Actualizar la configuración del agente",
         "La configuración se guarda y se devuelve sin error de serialización", "Aprobado"),
        ("CP-07", "Contactos", "Listar y crear contactos del negocio",
         "Las operaciones responden correctamente y solo con datos propios", "Aprobado"),
        ("CP-08", "Conversaciones", "Listar conversaciones del negocio",
         "Se devuelven las conversaciones del negocio autenticado", "Aprobado"),
        ("CP-09", "Métricas", "Obtener el resumen del panel y los gráficos",
         "Se devuelven métricas coherentes (mensajes, leads, llamadas)", "Aprobado"),
        ("CP-10", "WhatsApp", "Recibir un webhook con firma válida e inválida",
         "Se acepta la firma válida (200) y se rechaza la inválida (403)", "Aprobado"),
        ("CP-11", "Prueba gratuita", "Acceder con un periodo de prueba vencido",
         "Se bloquea el acceso con código 402 y se redirige a la mejora de plan", "Aprobado"),
        ("CP-12", "Cobros", "Confirmar un pago y suspender por falta de pago",
         "El servicio se reactiva al confirmar y se bloquea al suspender", "Aprobado"),
        ("CP-13", "Administración", "Crear, exportar y eliminar un negocio",
         "Las operaciones del panel de administración se ejecutan correctamente", "Aprobado"),
        ("CP-14", "Seguridad", "Superar el límite de intentos de inicio de sesión",
         "Se aplica el límite de tasa y se responde 429", "Aprobado"),
        ("CP-15", "Recuperación", "Solicitar y aprobar el restablecimiento de contraseña",
         "Se genera una contraseña nueva y la anterior deja de funcionar", "Aprobado"),
    ])

    h1(doc, "9. Resultados de la ejecución")
    para(doc,
         "La ejecución de la batería de pruebas automatizadas del backend arrojó un resultado "
         "satisfactorio. Se ejecutaron 543 pruebas con resultado favorable; 5 pruebas adicionales, "
         "marcadas como «de preparación» (readiness), quedan deseleccionadas de forma intencional "
         "porque requieren claves y configuración de servicios externos reales. En el frontend, las "
         "pruebas con Vitest se ejecutaron correctamente y la construcción de producción se completó "
         "sin errores.")
    kv_table(doc, [
        ("Pruebas de backend ejecutadas con éxito", "543"),
        ("Pruebas deseleccionadas (requieren claves reales)", "5"),
        ("Pruebas de frontend (Vitest)", "Ejecutadas correctamente"),
        ("Construcción de producción del frontend", "Sin errores"),
        ("Defectos críticos pendientes", "Ninguno"),
    ], headers=("Indicador", "Resultado"))

    h1(doc, "10. Defectos detectados y corregidos")
    para(doc,
         "Durante el proceso de pruebas se detectó y corrigió un defecto relevante: al actualizar la "
         "configuración del agente, una respuesta del servidor fallaba (error 500) debido a la lectura "
         "diferida de un campo de fecha fuera de su contexto de ejecución asíncrono. El defecto se "
         "corrigió forzando la actualización del registro tras la operación de guardado, y se añadió "
         "un caso de prueba para evitar su reaparición. Este hallazgo demuestra la utilidad de las "
         "pruebas automatizadas.")

    h1(doc, "11. Riesgos y mitigaciones")
    bullets(doc, [
        ("Dependencia de servicios externos",
         "se mitiga con simulaciones controladas y con degradación segura cuando faltan las claves."),
        ("Fuga de datos entre negocios",
         "se mitiga con el aislamiento de datos a nivel de aplicación, verificado por pruebas."),
        ("Regresiones al añadir funciones",
         "se mitigan ejecutando la batería completa de pruebas antes de cada despliegue."),
    ])

    h1(doc, "12. Conclusiones")
    para(doc,
         "El software AgentePro fue sometido a un conjunto amplio de pruebas unitarias, de "
         "integración, funcionales y de seguridad, con resultados satisfactorios y sin defectos "
         "críticos pendientes. El sistema cumple los criterios de salida definidos en este plan y se "
         "considera apto para su despliegue, sujeto a la configuración de las claves de los servicios "
         "externos en el entorno de producción.")
    para(doc,
         f"El presente plan fue elaborado por el equipo de desarrollo de AgentePro, estudiantes de la "
         f"{UNIVERSIDAD}, bajo la asesoría del docente {ASESOR}.")
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_______________________________").bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(AUTOR + "\n").bold = True
    p.add_run("Autor principal, titular y responsable de calidad\n")
    p.add_run(f"DNI N.º {DNI}\n")
    p.add_run(FECHA_FIRMA)

    doc.save(path)
    print("OK:", path)


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    build(os.path.join(here, "Plan_de_Pruebas_de_Software_AgentePro.docx"))
