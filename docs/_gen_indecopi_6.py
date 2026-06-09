# -*- coding: utf-8 -*-
"""Genera los documentos de registro ante INDECOPI (DDA) del software AgentePro,
YA RELLENADOS con los datos reales.

Autoria: obra en colaboracion. Autor principal y unico titular de los derechos
patrimoniales: Italo Eduardo Reyes Cordero. Coautores (estudiantes de la
Universidad Continental): Jack Joshua Bendezu Lagos y Dickmar Wilber Julca
Laureano. Docente asesor del proyecto: Maglioni Arana Caparachin.

Genera: Documentos 1 al 7 + un indice del expediente.
Ejecutar: python "_gen_indecopi_6.py"
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

AZUL = RGBColor(0x1D, 0x4E, 0xD8)
GRIS = RGBColor(0x55, 0x55, 0x55)
NEGRO = RGBColor(0x11, 0x11, 0x11)

# ── Autor principal y titular ───────────────────────────────────────────────
AUTOR = "Italo Eduardo Reyes Cordero"
CORREO = "italoreyescordero1@gmail.com"
DNI = "75220834"
DOMICILIO = "Jr. Grau N.º 419, distrito de Jauja, provincia de Jauja, departamento de Junín"
TELEFONO = "916085873"

# ── Coautores (estudiantes de la Universidad Continental) ───────────────────
COAUTOR2 = "Jack Joshua Bendezu Lagos"
DNI2 = "73940475"
COAUTOR3 = "Dickmar Wilber Julca Laureano"
DNI3 = "73086197"

# ── Docente asesor (rol de asesoría; no es autor legal de la obra) ──────────
ASESOR = "Maglioni Arana Caparachin"
DNI_ASESOR = "20038141"

UNIVERSIDAD = "Universidad Continental"
CIUDAD = "Huancayo, Perú"
CIUDAD_FIRMA = "Huancayo"
FECHA_FIRMA = "Huancayo, 5 de junio de 2026"
ANIO = "2026"
SOFTWARE = "AgentePro"

# Roles para reutilizar en varios documentos
ROL_AUTOR = ("Autor principal y titular de los derechos. Dirección técnica del "
             "proyecto; análisis y diseño de la arquitectura multiempresa; "
             "programación del servidor (backend en FastAPI); diseño y modelado de "
             "la base de datos; integración de los servicios de inteligencia "
             "artificial, mensajería y telefonía; panel de superadministración, "
             "seguridad, aislamiento de datos y despliegue.")
ROL_C2 = ("Coautor. Colaboró en el desarrollo de la interfaz de usuario (frontend "
          "en React y TypeScript): maquetación de páginas y componentes del panel, "
          "estilos y apoyo en la documentación de usuario y en las pruebas de "
          "interfaz.")
ROL_C3 = ("Coautor. Colaboró en el aseguramiento de la calidad: diseño y ejecución "
          "de casos de prueba, apoyo en el modelado y la validación de datos y en la "
          "elaboración de la documentación técnica del proyecto.")
ROL_ASESOR = ("Docente asesor del proyecto. Brindó orientación metodológica, "
              "supervisión académica y revisión del trabajo. No participó en la "
              "programación de la obra; su rol es de asesoría.")


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.font.color.rgb = NEGRO
    rpr = st.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:ascii"), "Calibri")
    rf.set(qn("w:hAnsi"), "Calibri")


def doc_header(doc, numero, titulo):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"DOCUMENTO {numero}"); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = GRIS
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo.upper()); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = AZUL
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Software: {SOFTWARE}  ·  Autor principal y titular: {AUTOR}  ·  {CIUDAD}, {ANIO}")
    r.font.size = Pt(10); r.font.color.rgb = GRIS
    doc.add_paragraph()


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text); r.font.color.rgb = AZUL; r.font.size = Pt(14); r.bold = True


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); r.font.color.rgb = NEGRO; r.font.size = Pt(12); r.bold = True


def para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullets(doc, items):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(it[0] + ": ").bold = True
            p.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def kv_table(doc, rows, headers=("Campo", "Detalle")):
    t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    h = t.rows[0].cells
    h[0].paragraphs[0].add_run(headers[0]).bold = True
    h[1].paragraphs[0].add_run(headers[1]).bold = True
    for a, b in rows:
        c = t.add_row().cells
        c[0].paragraphs[0].add_run(str(a)).bold = True
        c[1].paragraphs[0].add_run(str(b))
    doc.add_paragraph()
    return t


def contexto_autoria(doc):
    """Bloque comun: la obra es en colaboracion, en el marco academico."""
    para(doc, f"El software «{SOFTWARE}» es una obra de software desarrollada en "
              f"colaboración por estudiantes de la {UNIVERSIDAD}, bajo la asesoría "
              f"del docente {ASESOR}. La autoría principal y la titularidad de los "
              f"derechos corresponden a {AUTOR}, quien concibió, dirigió y programó "
              f"la mayor parte de la obra. Participaron como coautores, con aportes "
              f"complementarios, los estudiantes {COAUTOR2} y {COAUTOR3}.")


def firma_uno(doc, nombre, dni, cargo):
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("_______________________________").bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(nombre + "\n").bold = True
    p.add_run(cargo + "\n")
    p.add_run(f"DNI N.º {dni}")


def firmas(doc, signers, con_fecha=True):
    """signers: lista de (nombre, dni, cargo)."""
    for nombre, dni, cargo in signers:
        firma_uno(doc, nombre, dni, cargo)
    if con_fecha:
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(FECHA_FIRMA); r.italic = True; r.font.color.rgb = GRIS


SIGN_AUTORES = [
    (AUTOR, DNI, "Autor principal y titular de la obra"),
    (COAUTOR2, DNI2, "Coautor"),
    (COAUTOR3, DNI3, "Coautor"),
]


# ════════════════════ DOCUMENTO 1 — FICHA TÉCNICA ════════════════════
def doc1(path):
    d = Document(); base_style(d)
    doc_header(d, 1, "Ficha Técnica del Software")
    para(d, "La presente ficha técnica identifica de manera clara y técnica el programa de ordenador "
            "que se registra ante INDECOPI, describiendo su denominación, lenguajes, funcionalidad, "
            "fecha de creación, autoría y arquitectura.")

    h1(d, "1. Título del software")
    para(d, "El presente programa de ordenador se denomina «AgentePro». La denominación identifica de "
            "forma única a la plataforma y la distingue de otros sistemas, evitando nombres genéricos. "
            "AgentePro es el nombre comercial y técnico con el que se reconoce la obra.")

    h1(d, "2. Autoría")
    contexto_autoria(d)
    kv_table(d, [
        (AUTOR, f"DNI N.º {DNI} — Autor principal y titular de los derechos"),
        (COAUTOR2, f"DNI N.º {DNI2} — Coautor (estudiante de la {UNIVERSIDAD})"),
        (COAUTOR3, f"DNI N.º {DNI3} — Coautor (estudiante de la {UNIVERSIDAD})"),
        (ASESOR, f"DNI N.º {DNI_ASESOR} — Docente asesor del proyecto"),
    ], headers=("Persona", "Documento y rol"))

    h1(d, "3. Lenguajes de programación utilizados")
    para(d, "El software ha sido desarrollado principalmente en Python (para el servidor y la lógica "
            "de negocio), TypeScript y JavaScript con la biblioteca React (para la interfaz de "
            "usuario) y SQL sobre PostgreSQL (para la gestión de la base de datos). Se emplearon "
            "tecnologías complementarias como FastAPI y Socket.IO en el servidor, Vite y Tailwind CSS "
            "en la interfaz, y Docker para el empaquetado y despliegue. La elección de estos lenguajes "
            "responde a criterios de escalabilidad, seguridad, rendimiento y facilidad de "
            "mantenimiento.")
    kv_table(d, [
        ("Servidor (backend)", "Python (FastAPI, SQLAlchemy, Socket.IO, Celery)"),
        ("Interfaz (frontend)", "TypeScript y JavaScript (React, Vite, Tailwind CSS)"),
        ("Base de datos", "SQL sobre PostgreSQL (migraciones con Alembic)"),
        ("Infraestructura", "Docker, Docker Compose, servidor web Caddy, Redis"),
    ], headers=("Capa", "Lenguajes y tecnologías"))

    h1(d, "4. Funcionalidad principal")
    para(d, "AgentePro tiene como funcionalidad principal la automatización de la atención al cliente "
            "y la gestión comercial de un negocio mediante inteligencia artificial. El sistema atiende "
            "de forma automática y permanente los mensajes de WhatsApp e Instagram, contesta y realiza "
            "llamadas de voz, registra contactos y conversaciones, califica clientes potenciales, "
            "agenda citas, genera contenido para redes sociales, ejecuta seguimientos automáticos y "
            "muestra todas las métricas del negocio en un panel de control en tiempo real. Es una "
            "plataforma multiempresa: cada negocio cliente opera de forma aislada y segura.")

    h1(d, "5. Fecha de creación")
    para(d, "La primera versión funcional del software fue concluida en mayo de 2026, tras un proceso "
            "de desarrollo realizado durante el primer semestre de 2026, que incluyó las fases de "
            "análisis, diseño, programación, pruebas y documentación.")

    h1(d, "6. Arquitectura del sistema")
    para(d, "AgentePro se estructura bajo una arquitectura cliente-servidor. El servidor, desarrollado "
            "en Python con FastAPI, implementa la lógica de negocio, la seguridad, el aislamiento de "
            "datos entre empresas, la comunicación en tiempo real y la integración con servicios "
            "externos. La interfaz, desarrollada en React, ofrece el panel del negocio y el panel de "
            "superadministración. La base de datos PostgreSQL almacena la información de forma "
            "estructurada. Todo el sistema se despliega en contenedores Docker, con un servidor web "
            "Caddy que provee cifrado HTTPS automático. La arquitectura es modular y está preparada "
            "para funcionar en la nube.")

    h1(d, "7. Observaciones finales")
    para(d, "La presente ficha técnica constituye un resumen del software AgentePro y forma parte del "
            "paquete de documentos para su registro como programa de ordenador ante INDECOPI, conforme "
            "al Decreto Legislativo N.º 822.")
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 2 — EJEMPLAR DEL SOFTWARE ════════════════════
def doc2(path):
    d = Document(); base_style(d)
    doc_header(d, 2, "Ejemplar del Software")
    para(d, "El ejemplar del software constituye la evidencia técnica de la obra. Demuestra su "
            "existencia, funcionalidad y originalidad mediante el código fuente, un manual técnico y "
            "material gráfico de las interfaces.")

    h1(d, "1. Código fuente")
    para(d, "Se adjunta un archivo comprimido denominado «Codigo_Fuente_AgentePro.zip», que contiene "
            "una versión representativa del código fuente del software, organizada de la siguiente "
            "manera:")
    bullets(d, [
        ("Carpeta backend", "código del servidor en Python (FastAPI): modelos de datos, servicios, "
         "puntos de acceso de la interfaz de programación, webhooks, tareas en segundo plano y "
         "migraciones de la base de datos."),
        ("Carpeta frontend", "código de la interfaz en TypeScript y React: páginas, componentes, "
         "estados y utilidades."),
        ("Archivos de configuración y despliegue", "definiciones de contenedores Docker y del "
         "servidor web."),
        ("Documentación técnica", "archivos de apoyo en la carpeta de documentación."),
    ])
    para(d, "El archivo comprimido representa fielmente la autoría de los autores y está limpio de "
            "credenciales, contraseñas y datos sensibles, por lo que constituye un ejemplar suficiente "
            "para acreditar la existencia del software ante INDECOPI. Asimismo, en la «Plantilla de "
            "Registro General de Software» se incluye una muestra impresa y representativa del código "
            "fuente.")

    h1(d, "2. Manual técnico")
    para(d, "El manual técnico guía la instalación, configuración y uso del software. Se adjunta de "
            "forma completa en el documento «Manual de Usuario de AgentePro». A continuación se "
            "resumen sus apartados principales:")
    h2(d, "2.1. Requisitos del sistema")
    bullets(d, [
        "Servidor con sistema operativo Linux y Docker instalado.",
        "Nombre de dominio propio apuntando al servidor.",
        "Clave de la interfaz de programación de Anthropic para activar la inteligencia artificial.",
        "Cuenta de WhatsApp Business API por cada negocio, conectada mediante la nube de Meta o "
        "mediante un proveedor de mensajería como Twilio.",
        "Navegador web moderno (Google Chrome o Mozilla Firefox).",
    ])
    h2(d, "2.2. Instalación y despliegue")
    numbered(d, [
        "Instalar Docker y Docker Compose en el servidor.",
        "Apuntar el dominio (registro DNS de tipo A) a la dirección IP del servidor.",
        "Copiar el código fuente al servidor y crear el archivo de configuración «.env».",
        "Ejecutar el comando de despliegue: docker compose -f docker-compose.deploy.yml up -d --build.",
        "Esperar a que el sistema genere el certificado de seguridad (HTTPS) y verificar el dominio.",
    ])
    h2(d, "2.3. Uso básico")
    bullets(d, [
        "Iniciar sesión con correo y contraseña.",
        "Conectar el WhatsApp del negocio y configurar el agente de inteligencia artificial.",
        "Consultar conversaciones, llamadas, contactos y métricas en el panel.",
    ])
    h2(d, "2.4. Mantenimiento")
    bullets(d, [
        "Las copias de seguridad de la base de datos se realizan de forma automática y periódica.",
        "Las actualizaciones se aplican reconstruyendo los contenedores.",
    ])

    h1(d, "3. Capturas de pantalla")
    para(d, "Se adjuntan capturas de pantalla que muestran las principales interfaces del software:")
    bullets(d, [
        "Página pública de presentación, con la marca y los planes.",
        "Pantalla de inicio de sesión.",
        "Panel del negocio (dashboard): métricas, volumen de mensajes y embudo de clientes.",
        "Sección de conversaciones atendidas por el agente.",
        "Panel de superadministración: gestión de negocios, uso y cobros.",
    ])
    para(d, "Las capturas permiten visualizar la funcionalidad principal del sistema y constituyen "
            "evidencia gráfica de su existencia. Las imágenes no contienen datos personales reales.")

    h1(d, "4. Observaciones finales")
    para(d, "El ejemplar del software AgentePro constituye evidencia suficiente para acreditar su "
            "existencia y funcionalidad ante INDECOPI. Junto con la ficha técnica, la declaración "
            "jurada de autoría y los anexos técnicos, conforma el paquete completo de registro.")
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 3 — DECLARACIÓN JURADA DE AUTORÍA ════════════════════
def doc3(path):
    d = Document(); base_style(d)
    doc_header(d, 3, "Declaración Jurada de Autoría")
    para(d, "La presente Declaración Jurada de Autoría reconoce la autoría del software, declara su "
            "originalidad y deja constancia de la titularidad de los derechos.")

    h1(d, "1. Identificación de los declarantes")
    para(d, "Los abajo firmantes, en calidad de autores del software denominado «AgentePro», "
            "declaramos bajo juramento lo siguiente:")
    kv_table(d, [
        (AUTOR, f"DNI N.º {DNI} — Autor principal y titular de los derechos"),
        (COAUTOR2, f"DNI N.º {DNI2} — Coautor"),
        (COAUTOR3, f"DNI N.º {DNI3} — Coautor"),
    ], headers=("Declarante", "Documento y calidad"))

    h1(d, "2. Reconocimiento de autoría")
    para(d, "Declaramos ser los autores del software mencionado, desarrollado en los lenguajes de "
            "programación Python, TypeScript, JavaScript y SQL, cuya funcionalidad principal consiste "
            "en automatizar la atención al cliente y la gestión comercial de un negocio mediante "
            "inteligencia artificial. La obra fue concebida y desarrollada como un trabajo en "
            "colaboración en el marco académico de la Universidad Continental, bajo la asesoría del "
            "docente Maglioni Arana Caparachin.")

    h1(d, "3. Autoría y aportes")
    para(d, "La obra es una obra en colaboración. Los aportes de cada autor son los siguientes:")
    bullets(d, [
        (AUTOR, ROL_AUTOR),
        (COAUTOR2, ROL_C2),
        (COAUTOR3, ROL_C3),
    ])
    para(d, f"El autor principal, {AUTOR}, concibió, dirigió y desarrolló la mayor parte de la obra. "
            f"El docente {ASESOR} (DNI N.º {DNI_ASESOR}) participó únicamente como asesor del proyecto, "
            f"sin intervenir en la programación, por lo que no se le considera autor de la obra.")

    h1(d, "4. Declaración de originalidad")
    para(d, "Declaramos que la obra es original, producto de nuestro trabajo creativo, y que no "
            "infringe derechos de terceros. El software no ha sido copiado ni adaptado de programas "
            "ajenos sin autorización, y constituye una creación independiente. El uso de bibliotecas y "
            "marcos de trabajo de terceros se realiza conforme a sus respectivas licencias de código "
            "abierto, sin que ello afecte la originalidad de la obra desarrollada.")

    h1(d, "5. Titularidad de los derechos")
    para(d, f"Los coautores {COAUTOR2} y {COAUTOR3} ceden a favor de {AUTOR} la totalidad de sus "
            f"derechos patrimoniales sobre la obra, de manera que la titularidad de dichos derechos "
            f"recae de forma única y exclusiva en {AUTOR}, conservando cada coautor sus derechos "
            f"morales como tal. En consecuencia, se solicita y autoriza el registro de la obra "
            f"consignando a {AUTOR} como autor principal y titular de los derechos patrimoniales, de "
            f"conformidad con el Decreto Legislativo N.º 822, Ley sobre el Derecho de Autor. El detalle "
            f"de la cesión consta en el documento «Cesión de Derechos Patrimoniales».")

    h1(d, "6. Cláusula de juramento")
    para(d, "Declaramos bajo juramento la veracidad de lo expuesto en el presente documento, "
            "comprometiéndonos a responder legalmente en caso de falsedad.")

    h1(d, "7. Firmas y fecha")
    firmas(d, SIGN_AUTORES)
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 4 — LISTA DE AUTORES Y ROLES ════════════════════
def doc4(path):
    d = Document(); base_style(d)
    doc_header(d, 4, "Lista de Autores y Roles")
    para(d, "El presente documento identifica a los autores del software y detalla los roles que "
            "desempeñaron durante su creación, así como la participación del docente asesor.")

    h1(d, "1. Identificación del software")
    contexto_autoria(d)

    h1(d, "2. Tabla de autores y roles")
    t = d.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
    heads = ("Nombre completo / DNI", "Rol en el proyecto", "Componentes desarrollados", "Periodo")
    for i, h in enumerate(heads):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True

    filas = [
        (f"{AUTOR}\nDNI N.º {DNI}",
         "Autor principal y titular. Desarrollador integral (full-stack) y director técnico del "
         "proyecto.",
         "Arquitectura multiempresa, servidor (backend en FastAPI), base de datos y migraciones, "
         "integraciones de IA, mensajería y voz, panel de superadministración, seguridad, aislamiento "
         "de datos, despliegue y batería de pruebas.",
         "Primer semestre de 2026"),
        (f"{COAUTOR2}\nDNI N.º {DNI2}",
         "Coautor (estudiante de la Universidad Continental).",
         "Interfaz de usuario (frontend en React y TypeScript): páginas y componentes del panel, "
         "estilos, apoyo en documentación de usuario y pruebas de interfaz.",
         "Primer semestre de 2026"),
        (f"{COAUTOR3}\nDNI N.º {DNI3}",
         "Coautor (estudiante de la Universidad Continental).",
         "Aseguramiento de la calidad: diseño y ejecución de casos de prueba, apoyo en el modelado y "
         "validación de datos y en la documentación técnica.",
         "Primer semestre de 2026"),
    ]
    for a, b, c, e in filas:
        row = t.add_row().cells
        row[0].paragraphs[0].add_run(a)
        row[1].paragraphs[0].add_run(b)
        row[2].paragraphs[0].add_run(c)
        row[3].paragraphs[0].add_run(e)
    d.add_paragraph()

    h1(d, "3. Docente asesor")
    kv_table(d, [
        ("Nombre", ASESOR),
        ("Documento de identidad", f"DNI N.º {DNI_ASESOR}"),
        ("Institución", UNIVERSIDAD),
        ("Rol", ROL_ASESOR),
    ])
    para(d, "Se deja constancia de que el docente asesor no es autor de la obra; su participación se "
            "limitó a la orientación y supervisión académica del proyecto.")

    h1(d, "4. Descripción de los aportes")
    para(d, "El autor principal diseñó la arquitectura multiempresa con aislamiento de datos, programó "
            "la mayor parte de la lógica de negocio del servidor, modeló la base de datos, integró los "
            "servicios de inteligencia artificial, mensajería y telefonía, e implementó el panel de "
            "superadministración y la batería de pruebas automatizadas. Los coautores colaboraron, "
            "respectivamente, en el desarrollo de la interfaz de usuario y en el aseguramiento de la "
            "calidad y la documentación, con aportes complementarios al trabajo del autor principal.")

    h1(d, "5. Validación")
    para(d, "Los autores validamos la presente lista de autores y roles, que forma parte del paquete "
            "de registro ante INDECOPI.")

    h1(d, "6. Firmas y fecha")
    firmas(d, SIGN_AUTORES)
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 5 — TITULARIDAD (PERSONA NATURAL) ════════════════════
def doc5(path):
    d = Document(); base_style(d)
    doc_header(d, 5, "Titularidad de la Obra (Persona Natural)")
    para(d, "El presente documento acredita la titularidad del software. El titular de los derechos "
            "patrimoniales es una persona natural: el autor principal de la obra. Los coautores "
            "conservan sus derechos morales y han cedido sus derechos patrimoniales al titular.")

    h1(d, "1. Naturaleza de la titularidad")
    para(d, f"El software «{SOFTWARE}» es una obra en colaboración cuyos derechos patrimoniales recaen, "
            f"de forma íntegra y exclusiva, en una persona natural: {AUTOR}, autor principal de la "
            f"obra. Los coautores {COAUTOR2} y {COAUTOR3} han cedido a su favor la totalidad de sus "
            f"derechos patrimoniales, conservando sus respectivos derechos morales. No existe persona "
            f"jurídica (empresa) titular de los derechos ni representante legal que la gestione; el "
            f"titular actúa por derecho propio.")

    h1(d, "2. Identificación del titular")
    kv_table(d, [
        ("Titular y autor principal", AUTOR),
        ("Tipo de titular", "Persona natural"),
        ("Documento de identidad", f"DNI N.º {DNI}"),
        ("Domicilio", DOMICILIO),
        ("Nacionalidad", "Peruana"),
        ("Correo electrónico", CORREO),
        ("Teléfono", TELEFONO),
    ])

    h1(d, "3. Coautores y cesión de derechos")
    para(d, "Los coautores de la obra, estudiantes de la Universidad Continental, han cedido sus "
            "derechos patrimoniales al titular antes identificado, conforme al documento «Cesión de "
            "Derechos Patrimoniales». Conservan, en su calidad de autores, los derechos morales que la "
            "ley les reconoce.")
    kv_table(d, [
        (COAUTOR2, f"DNI N.º {DNI2} — Coautor (cede derechos patrimoniales)"),
        (COAUTOR3, f"DNI N.º {DNI3} — Coautor (cede derechos patrimoniales)"),
    ], headers=("Coautor", "Documento y situación"))

    h1(d, "4. Declaración de titularidad")
    para(d, f"{AUTOR} declara ser el titular de los derechos patrimoniales del software «{SOFTWARE}» y "
            f"autor principal de la obra. En ejercicio de sus derechos patrimoniales, se reserva las "
            f"facultades de reproducción, distribución, comunicación pública y transformación de la "
            f"obra, conforme al Decreto Legislativo N.º 822, Ley sobre el Derecho de Autor.")

    h1(d, "5. Facultades para el registro")
    para(d, "En su condición de autor principal y titular, se encuentra plenamente facultado para "
            "realizar todos los actos necesarios para el registro de la obra ante INDECOPI, incluyendo "
            "la presentación de la solicitud, los anexos, el ejemplar del software y la declaración "
            "jurada de autoría, así como para suscribir, en el futuro, contratos de licencia o cesión "
            "de derechos sobre la obra.")

    h1(d, "6. Nota sobre una eventual empresa titular")
    para(d, "En caso de que, en el futuro, el titular constituya una empresa (por ejemplo, una "
            "E.I.R.L. o una S.A.C.) y desee que la titularidad de los derechos patrimoniales "
            "corresponda a dicha empresa, deberá formalizarse una cesión de derechos del titular a la "
            "empresa y actualizarse el registro. A la fecha del presente documento, la titularidad "
            "corresponde a la persona natural antes identificada.")

    h1(d, "7. Cláusula de responsabilidad")
    para(d, "El titular asume plena responsabilidad por la veracidad de la información contenida en "
            "el presente documento y por la autenticidad de los anexos que se adjuntan, "
            "comprometiéndose a responder conforme a la legislación vigente en caso de falsedad.")

    h1(d, "8. Firma y fecha")
    firmas(d, [(AUTOR, DNI, "Autor principal y titular de la obra (persona natural)")])
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 6 — ANEXOS TÉCNICOS ════════════════════
def doc6(path):
    d = Document(); base_style(d)
    doc_header(d, 6, "Anexos Técnicos")
    para(d, "Los anexos técnicos constituyen material complementario que respalda la descripción y el "
            "ejemplar del software, demostrando su estructura interna, su funcionamiento, sus "
            "interfaces y los procesos de calidad aplicados.")

    h1(d, "1. Flujo principal del sistema")
    para(d, "El flujo principal de AgentePro, para la atención automática de un cliente, comprende las "
            "siguientes etapas:")
    numbered(d, [
        "Un cliente envía un mensaje al WhatsApp o Instagram del negocio.",
        "La plataforma recibe el evento mediante un webhook y lo asocia al negocio correspondiente.",
        "El motor de inteligencia artificial genera una respuesta según la configuración del agente.",
        "La respuesta se envía al cliente y la conversación se registra en la base de datos.",
        "El sistema califica al cliente potencial, agenda citas y programa seguimientos según "
        "corresponda.",
        "El negocio visualiza la actividad y las métricas en su panel de control en tiempo real.",
    ])
    para(d, "Se recomienda adjuntar un diagrama de flujo (por ejemplo, en notación BPMN o UML) que "
            "represente gráficamente este proceso.")

    h1(d, "2. Arquitectura del software")
    para(d, "AgentePro se estructura bajo una arquitectura cliente-servidor, con los siguientes "
            "componentes:")
    bullets(d, [
        ("Servidor (backend)", "desarrollado en Python con FastAPI; gestiona la lógica de negocio, la "
         "seguridad, el aislamiento de datos entre empresas y las integraciones externas; expone una "
         "interfaz de programación y procesa eventos mediante webhooks; usa Socket.IO para la "
         "comunicación en tiempo real y Celery para las tareas en segundo plano."),
        ("Interfaz (frontend)", "desarrollada en TypeScript con React; comprende la página pública, el "
         "panel del negocio y el panel de superadministración."),
        ("Base de datos", "PostgreSQL, con un esquema versionado mediante migraciones."),
        ("Infraestructura", "contenedores Docker, servidor web Caddy con HTTPS automático, memoria "
         "caché Redis y un servicio de copias de seguridad automáticas."),
    ])
    para(d, "La arquitectura es modular y está preparada para funcionar en la nube. Se recomienda "
            "adjuntar un diagrama de arquitectura como anexo gráfico.")

    h1(d, "3. Capturas de pantalla de módulos clave")
    bullets(d, [
        "Página pública de presentación con la marca y los planes.",
        "Pantalla de inicio de sesión.",
        "Panel del negocio: métricas, volumen de mensajes y embudo de clientes potenciales.",
        "Sección de conversaciones y de configuración del agente de inteligencia artificial.",
        "Panel de superadministración: gestión de negocios, uso y consumo, y cobros.",
    ])

    h1(d, "4. Resumen de pruebas de calidad (QA)")
    para(d, "El software cuenta con una batería de pruebas automatizadas que verifica su correcto "
            "funcionamiento. El resumen de resultados es el siguiente:")
    kv_table(d, [
        ("Pruebas de servidor ejecutadas con éxito", "543"),
        ("Pruebas deseleccionadas (requieren claves de servicios reales)", "5"),
        ("Tipos de prueba", "Unitarias, de integración (HTTP), funcionales y de seguridad"),
        ("Pruebas de interfaz", "Ejecutadas correctamente con Vitest"),
        ("Construcción de producción del frontend", "Sin errores"),
        ("Defectos críticos pendientes", "Ninguno"),
    ], headers=("Indicador", "Resultado"))
    para(d, "El detalle completo se encuentra en el documento «Plan de Pruebas de Software de "
            "AgentePro».")

    h1(d, "5. Documentación complementaria")
    bullets(d, [
        "Manual de Usuario de AgentePro (instalación, configuración y operación).",
        "Memoria Descriptiva de la obra.",
        "Plan de Pruebas de Software.",
        "Ficha Técnica del Software (Documento 1).",
        "Cesión de Derechos Patrimoniales (Documento 7).",
    ])

    h1(d, "6. Observaciones finales")
    para(d, "Los anexos técnicos presentados constituyen evidencia adicional que respalda la "
            "originalidad y funcionalidad del software AgentePro, y complementan el paquete de "
            "documentos para su registro ante INDECOPI.")
    d.save(path); print("OK:", path)


# ════════════════════ DOCUMENTO 7 — CESIÓN DE DERECHOS PATRIMONIALES ════════════════════
def doc7(path):
    d = Document(); base_style(d)
    doc_header(d, 7, "Cesión de Derechos Patrimoniales")
    para(d, "El presente documento formaliza la cesión de los derechos patrimoniales de los coautores "
            "del software «AgentePro» a favor del autor principal, a fin de que la titularidad de "
            "dichos derechos recaiga de forma única y exclusiva en una sola persona, para su registro "
            "ante INDECOPI.")

    h1(d, "1. Partes")
    kv_table(d, [
        ("Cesionario (titular)", f"{AUTOR} — DNI N.º {DNI}"),
        ("Cedente 1 (coautor)", f"{COAUTOR2} — DNI N.º {DNI2}"),
        ("Cedente 2 (coautor)", f"{COAUTOR3} — DNI N.º {DNI3}"),
    ], headers=("Calidad", "Identificación"))

    h1(d, "2. Antecedente")
    para(d, "El software «AgentePro» fue desarrollado como una obra en colaboración por los autores "
            "antes identificados, estudiantes de la Universidad Continental, bajo la asesoría del "
            "docente Maglioni Arana Caparachin. El autor principal concibió, dirigió y desarrolló la "
            "mayor parte de la obra; los coautores realizaron aportes complementarios.")

    h1(d, "3. Objeto de la cesión")
    para(d, f"Los cedentes, {COAUTOR2} y {COAUTOR3}, ceden y transfieren a favor de {AUTOR} la "
            f"totalidad de los derechos patrimoniales que les corresponden como coautores del software "
            f"«AgentePro», incluyendo, de manera enunciativa y no limitativa, las facultades de "
            f"reproducción, distribución, comunicación pública, traducción, adaptación, modificación y "
            f"transformación de la obra, así como cualquier otra forma de explotación, en todo el mundo "
            f"y por todo el plazo de protección que reconoce la ley.")

    h1(d, "4. Carácter de la cesión")
    para(d, "La cesión se realiza con carácter exclusivo, ilimitado y a título gratuito. En "
            "consecuencia, el cesionario queda como único titular de los derechos patrimoniales de la "
            "obra y puede ejercerlos y disponer de ellos sin necesidad de autorización adicional de "
            "los cedentes.")

    h1(d, "5. Derechos morales")
    para(d, "La presente cesión recae únicamente sobre los derechos patrimoniales. Los cedentes "
            "conservan, en su calidad de coautores, los derechos morales que la ley les reconoce, en "
            "particular el derecho a la paternidad de la obra, de conformidad con el Decreto "
            "Legislativo N.º 822, Ley sobre el Derecho de Autor.")

    h1(d, "6. Declaración")
    para(d, "Las partes declaran que suscriben el presente documento de forma libre y voluntaria, y "
            "que la información consignada es veraz, asumiendo responsabilidad conforme a la "
            "legislación vigente en caso de falsedad.")

    h1(d, "7. Firmas y fecha")
    firmas(d, [
        (AUTOR, DNI, "Cesionario · Autor principal y titular"),
        (COAUTOR2, DNI2, "Cedente · Coautor"),
        (COAUTOR3, DNI3, "Cedente · Coautor"),
    ])
    d.save(path); print("OK:", path)


# ════════════════════ ÍNDICE DEL EXPEDIENTE ════════════════════
def indice(path):
    d = Document(); base_style(d)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EXPEDIENTE DE REGISTRO DE PROGRAMA DE ORDENADOR")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = AZUL
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Software: {SOFTWARE}  ·  INDECOPI — Dirección de Derecho de Autor")
    r.font.size = Pt(11); r.font.color.rgb = GRIS
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(FECHA_FIRMA); r.italic = True; r.font.color.rgb = GRIS
    d.add_paragraph()

    h1(d, "Datos generales")
    kv_table(d, [
        ("Título de la obra", SOFTWARE),
        ("Tipo de obra", "Programa de ordenador (software)"),
        ("Autor principal y titular", f"{AUTOR} — DNI N.º {DNI}"),
        ("Coautores", f"{COAUTOR2} (DNI N.º {DNI2}); {COAUTOR3} (DNI N.º {DNI3})"),
        ("Docente asesor", f"{ASESOR} — DNI N.º {DNI_ASESOR}"),
        ("Institución", UNIVERSIDAD),
        ("Estado de la obra", "Inédita (no publicada)"),
        ("Año de creación", ANIO),
        ("Lugar", CIUDAD),
    ])

    h1(d, "Documentos que conforman el expediente")
    items = [
        "Formulario de solicitud de registro de programa de ordenador (INDECOPI), rellenado.",
        "Documento 1 — Ficha Técnica del Software.",
        "Documento 2 — Ejemplar del Software.",
        "Documento 3 — Declaración Jurada de Autoría.",
        "Documento 4 — Lista de Autores y Roles.",
        "Documento 5 — Titularidad de la Obra (persona natural).",
        "Documento 6 — Anexos Técnicos.",
        "Documento 7 — Cesión de Derechos Patrimoniales.",
        "Memoria Descriptiva de la obra.",
        "Manual de Usuario de AgentePro.",
        "Plan de Pruebas de Software.",
        "Plantilla de Registro General de Software (ejemplar de código y descripción técnica).",
        "Código fuente representativo (Codigo_Fuente_AgentePro.zip).",
    ]
    numbered(d, items)

    h1(d, "Nota")
    para(d, "El derecho de autor sobre la obra nace con su creación. El registro ante INDECOPI tiene "
            "por objeto constituir un medio de prueba de la titularidad y de la fecha cierta de la "
            "creación. Los datos personales de los coautores (fecha de nacimiento, dirección y "
            "teléfono) que el formulario solicite deberán ser completados por cada coautor antes de la "
            "presentación.")
    d.save(path); print("OK:", path)


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    p = lambda n: os.path.join(here, n)
    doc1(p("Documento 1 - Ficha técnica de software - AgentePro.docx"))
    doc2(p("Documento 2 - Ejemplar del software - AgentePro.docx"))
    doc3(p("Documento 3 - Declaración jurada de autoría - AgentePro.docx"))
    doc4(p("Documento 4 - Lista de autores y roles - AgentePro.docx"))
    doc5(p("Documento 5 - Titularidad de la obra - AgentePro.docx"))
    doc6(p("Documento 6 - Anexos técnicos - AgentePro.docx"))
    doc7(p("Documento 7 - Cesión de derechos patrimoniales - AgentePro.docx"))
    indice(p("Documento 0 - Índice del expediente - AgentePro.docx"))
