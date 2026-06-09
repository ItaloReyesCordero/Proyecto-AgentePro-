# -*- coding: utf-8 -*-
"""Rellena el formulario oficial de INDECOPI (Solicitud de registro de programa
de ordenador) con los datos reales de AgentePro, conservando toda la estructura
y el texto original de la plantilla. Genera una COPIA rellenada; la plantilla
en blanco se conserva intacta.
Ejecutar: python "_gen_formulario.py"
"""
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import os

AZUL = RGBColor(0x1D, 0x4E, 0xD8)

# ── Datos ───────────────────────────────────────────────────────────────────
AUTOR = "Italo Eduardo Reyes Cordero"
DNI = "75220834"
COAUTOR2 = "Jack Joshua Bendezu Lagos"
DNI2 = "73940475"
COAUTOR3 = "Dickmar Wilber Julca Laureano"
DNI3 = "73086197"
CORREO = "italoreyescordero1@gmail.com"
TELEFONO = "916085873"


def app(p, text, bold=True):
    """Anexa un run (resaltado) al final del parrafo, conservando el texto previo."""
    r = p.add_run(text)
    r.bold = bold
    if bold:
        r.font.color.rgb = AZUL
    return r


def reset(p, segments):
    """Reescribe el parrafo: segments = lista de (texto, bold)."""
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    for text, bold in segments:
        r = p.add_run(text)
        r.bold = bold
        if bold:
            r.font.color.rgb = AZUL


def insert_after(p, text):
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    r = np.add_run(text)
    r.bold = True
    r.font.color.rgb = AZUL
    return np


def fill(src, dst):
    d = Document(src)
    T0, T1, T2, T3 = d.tables[0], d.tables[1], d.tables[2], d.tables[3]

    # ── DATOS DEL SOLICITANTE (Italo, persona natural) ──
    c = T0.rows[0].cells[0]
    insert_after(c.paragraphs[0], AUTOR.upper())
    c = T0.rows[1].cells[0]
    reset(c.paragraphs[1], [("Sexo:   M  [ X ]      F  [   ]", True)])
    c = T0.rows[2].cells[0]
    reset(c.paragraphs[0], [("DNI: ", False), (DNI, True)])

    # ── DATOS DE LA OBRA ──
    app(T1.rows[0].cells[0].paragraphs[0], ":   AgentePro")
    reset(T1.rows[1].cells[0].paragraphs[1],
          [("¿Cuándo se terminó la obra?   ", False), ("31 / 05 / 2026", True),
           ("        País de origen:  ", False), ("Perú", True)])
    # ¿La obra se publicó?  -> No (inédita)
    reset(T1.rows[2].cells[0].paragraphs[1], [("Si  [   ]        No  [ X ]", True)])
    # ¿La obra es derivada? -> No
    reset(T1.rows[3].cells[0].paragraphs[1], [("Si  [   ]        No  [ X ]", True)])

    # ── DATOS DEL AUTOR O AUTORES ──
    autores = [
        # (cell_index, nombre, dni, sexo_M, completo)
        (0, AUTOR, DNI, True, True),
        (1, COAUTOR2, DNI2, True, False),
        (2, COAUTOR3, DNI3, True, False),
    ]
    for idx, nombre, dni, sexoM, completo in autores:
        cell = T2.rows[idx].cells[0]
        ps = cell.paragraphs
        sexo = "M [ X ]   F [   ]" if sexoM else "M [   ]   F [   ]"
        reset(ps[1], [("Nombres completos (conforme aparece en su documento de identidad)        Sexo:  ", False),
                      (sexo, True)])
        reset(ps[2], [(nombre, True)])
        reset(ps[6], [("DNI — N.º ", False), (dni, True)])
        if completo:
            reset(ps[8], [("País de nacimiento:  ", False), ("Perú", True),
                          ("      Fecha de nacimiento:  ____/____/______", False)])
            reset(ps[10], [("Jr. Grau N.º 419", True)])
            reset(ps[11], [("Distrito: ", False), ("Jauja", True),
                           ("      Provincia: ", False), ("Jauja", True),
                           ("      Departamento: ", False), ("Junín", True),
                           ("      Teléfono: ", False), (TELEFONO, True)])
        else:
            reset(ps[8], [("País de nacimiento:  ", False), ("Perú", True),
                          ("      Fecha de nacimiento:  ____/____/______", False)])
            reset(ps[10], [("____________________________  (a completar por el coautor)", False)])
            reset(ps[11], [("Distrito: __________   Provincia: __________   "
                            "Departamento: __________   Teléfono: __________", False)])

    # ── DOMICILIO PARA NOTIFICACIONES (titular Italo) ──
    c = T3.rows[0].cells[0]
    reset(c.paragraphs[1], [("Jr. Grau N.º 419", True)])
    reset(c.paragraphs[3], [("Distrito: ", False), ("Jauja", True),
                            ("      Provincia: ", False), ("Jauja", True),
                            ("      Departamento: ", False), ("Junín", True),
                            ("      Teléfono: ", False), (TELEFONO, True)])
    app(T3.rows[2].cells[0].paragraphs[0], ":   " + CORREO)

    # ── DATOS DEL PRODUCTOR Y TITULAR + Firma (parrafos del cuerpo) ──
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith("Se presume, salvo pacto en contrario"):
            insert_after(
                p,
                f"Productor y titular: {AUTOR} — DNI N.º {DNI}. Titular único de los "
                f"derechos patrimoniales. Los coautores cedieron sus derechos "
                f"patrimoniales a su favor (ver documento de cesión de derechos).")
        elif t == "Firma del solicitante":
            insert_after(p, f"{AUTOR} — DNI N.º {DNI}")

    d.save(dst)
    print("OK:", dst)


if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))
    src = os.path.join(here, "indecopi formiulario derecho de software indecopi.docx")
    dst = os.path.join(here, "Formulario INDECOPI - AgentePro (rellenado).docx")
    fill(src, dst)
